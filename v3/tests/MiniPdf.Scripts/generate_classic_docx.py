"""
Generate 30 classic DOCX test files for the MiniPdf benchmark suite.

Each file tests a different Word document feature, from simple paragraphs
to tables, images, lists, headings, and mixed content.

Prerequisites:
    pip install python-docx Pillow

Usage:
    python generate_classic_docx.py            # output to ./output_docx/
    python generate_classic_docx.py --outdir /path/to/dir
"""

import argparse
import io
import os
import struct
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

OUTPUT_DIR = Path(__file__).parent / "output_docx"


def _create_test_png(width=120, height=80, color=(70, 130, 180)):
    """Create a minimal PNG in memory for embedding."""
    if HAS_PIL:
        img = PILImage.new("RGB", (width, height), color)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf
    # Fallback: build a minimal 1x1 PNG manually
    return _minimal_png(color)


def _minimal_png(color):
    """Build a minimal valid 1×1 PNG."""
    import zlib
    buf = io.BytesIO()
    buf.write(b"\x89PNG\r\n\x1a\n")
    # IHDR
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    _write_chunk(buf, b"IHDR", ihdr_data)
    # IDAT
    raw = b"\x00" + bytes(color)
    compressed = zlib.compress(raw)
    _write_chunk(buf, b"IDAT", compressed)
    # IEND
    _write_chunk(buf, b"IEND", b"")
    buf.seek(0)
    return buf


def _write_chunk(buf, chunk_type, data):
    import zlib
    buf.write(struct.pack(">I", len(data)))
    buf.write(chunk_type)
    buf.write(data)
    buf.write(struct.pack(">I", zlib.crc32(chunk_type + data) & 0xFFFFFFFF))


# ── Classic docx generators (1–60) ──────────────────────────────────────


def classic01_single_paragraph(path):
    """Single plain-text paragraph."""
    doc = Document()
    doc.add_paragraph("Hello, World! This is a simple single paragraph document created for benchmarking MiniPdf DOCX-to-PDF conversion.")
    doc.save(path)


def classic02_multiple_paragraphs(path):
    """Multiple plain-text paragraphs with implicit spacing."""
    doc = Document()
    for i in range(1, 6):
        doc.add_paragraph(f"This is paragraph {i}. It contains some sample text to test how MiniPdf handles multiple consecutive paragraphs with default spacing.")
    doc.save(path)


def classic03_headings(path):
    """Heading levels 1 through 4."""
    doc = Document()
    doc.add_heading("Heading Level 1", level=1)
    doc.add_paragraph("Content under heading 1.")
    doc.add_heading("Heading Level 2", level=2)
    doc.add_paragraph("Content under heading 2.")
    doc.add_heading("Heading Level 3", level=3)
    doc.add_paragraph("Content under heading 3.")
    doc.add_heading("Heading Level 4", level=4)
    doc.add_paragraph("Content under heading 4.")
    doc.save(path)


def classic04_bold_italic(path):
    """Bold, italic, and bold-italic text."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Normal text. ")
    run_b = p.add_run("Bold text. ")
    run_b.bold = True
    run_i = p.add_run("Italic text. ")
    run_i.italic = True
    run_bi = p.add_run("Bold and italic text.")
    run_bi.bold = True
    run_bi.italic = True
    doc.save(path)


def classic05_font_sizes(path):
    """Various font sizes in one document."""
    doc = Document()
    for size in [8, 10, 12, 14, 18, 24, 36]:
        p = doc.add_paragraph()
        run = p.add_run(f"This text is {size}pt.")
        run.font.size = Pt(size)
    doc.save(path)


def classic06_font_colors(path):
    """Colored text runs."""
    doc = Document()
    colors = [
        ("Red text", RGBColor(255, 0, 0)),
        ("Green text", RGBColor(0, 128, 0)),
        ("Blue text", RGBColor(0, 0, 255)),
        ("Orange text", RGBColor(255, 165, 0)),
        ("Purple text", RGBColor(128, 0, 128)),
    ]
    for text, color in colors:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.color.rgb = color
    doc.save(path)


def classic07_alignment(path):
    """Paragraph alignment: left, center, right, justify."""
    doc = Document()
    lorem = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua."

    p = doc.add_paragraph(lorem)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_run = p.runs[0] if p.runs else p.add_run("")
    doc.add_paragraph()

    p2 = doc.add_paragraph(lorem)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = doc.add_paragraph(lorem)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p4 = doc.add_paragraph(lorem)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(path)


def classic08_bullet_list(path):
    """Bullet (unordered) list."""
    doc = Document()
    doc.add_heading("Shopping List", level=2)
    items = ["Apples", "Bananas", "Cherries", "Dates", "Elderberries"]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)


def classic09_numbered_list(path):
    """Numbered (ordered) list."""
    doc = Document()
    doc.add_heading("Steps to Success", level=2)
    steps = ["Define the goal", "Research the topic", "Create a plan", "Execute the plan", "Review results"]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.save(path)


def classic10_simple_table(path):
    """Simple 3x4 table with header row."""
    doc = Document()
    doc.add_heading("Employee Directory", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Name", "Department", "Email"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Alice Johnson", "Engineering", "alice@example.com"],
        ["Bob Smith", "Marketing", "bob@example.com"],
        ["Carol Williams", "Finance", "carol@example.com"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = val
    doc.save(path)


def classic11_table_with_shading(path):
    """Table with alternating row shading."""
    doc = Document()
    doc.add_heading("Quarterly Sales", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Quarter", "Revenue", "Expenses", "Profit"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "4472C4")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    data = [
        ["Q1 2025", "$120,000", "$80,000", "$40,000"],
        ["Q2 2025", "$135,000", "$85,000", "$50,000"],
        ["Q3 2025", "$150,000", "$90,000", "$60,000"],
        ["Q4 2025", "$160,000", "$95,000", "$65,000"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            if ri % 2 == 0:
                _set_cell_shading(cell, "D9E2F3")
    doc.save(path)


def classic12_merged_cells_table(path):
    """Table with merged cells."""
    doc = Document()
    doc.add_heading("Schedule", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Time"
    table.rows[0].cells[1].text = "Monday"
    table.rows[0].cells[2].text = "Tuesday"
    table.rows[1].cells[0].text = "9:00 AM"
    merged = table.rows[1].cells[1].merge(table.rows[1].cells[2])
    merged.text = "Team Meeting"
    table.rows[2].cells[0].text = "10:00 AM"
    table.rows[2].cells[1].text = "Code Review"
    table.rows[2].cells[2].text = "Design Review"
    table.rows[3].cells[0].text = "2:00 PM"
    table.rows[3].cells[1].text = "Sprint Planning"
    table.rows[3].cells[2].text = "Retrospective"
    doc.save(path)


def classic13_long_document(path):
    """A longer document that spans multiple pages."""
    doc = Document()
    doc.add_heading("Project Report", level=1)
    doc.add_paragraph("This document is designed to span multiple pages to test pagination in MiniPdf.")
    for i in range(1, 16):
        doc.add_heading(f"Section {i}", level=2)
        doc.add_paragraph(
            f"This is section {i} of the report. It contains detailed analysis of the topic at hand. "
            "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs. "
            "How vexingly quick daft zebras jump. The five boxing wizards jump quickly. "
            "Sphinx of black quartz, judge my vow." * 2
        )
    doc.save(path)


def classic14_mixed_content(path):
    """A document with headings, paragraphs, and a table mixed together."""
    doc = Document()
    doc.add_heading("Monthly Report", level=1)
    doc.add_paragraph("This report summarizes the key metrics for the month of January 2026.")

    doc.add_heading("Revenue Summary", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Category"
    table.rows[0].cells[1].text = "Amount"
    for i, (cat, amt) in enumerate([("Product Sales", "$85,000"), ("Services", "$42,000"), ("Subscriptions", "$28,000")]):
        table.rows[i + 1].cells[0].text = cat
        table.rows[i + 1].cells[1].text = amt

    doc.add_heading("Key Observations", level=2)
    doc.add_paragraph("Product sales increased by 15% compared to the previous quarter.")
    doc.add_paragraph("Service revenue remained stable with a slight upward trend.")

    doc.add_heading("Action Items", level=2)
    for item in ["Expand marketing campaign", "Hire two additional engineers", "Launch new subscription tier"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)


def classic15_indentation(path):
    """Paragraphs with various indentation levels."""
    doc = Document()
    doc.add_heading("Indentation Test", level=2)
    levels = [0, 36, 72, 108, 144]
    for pts in levels:
        p = doc.add_paragraph(f"This paragraph is indented by {pts} points from the left margin.")
        p.paragraph_format.left_indent = Pt(pts)
    # First-line indent
    p2 = doc.add_paragraph("This paragraph has a first-line indent of 36 points. The remaining lines wrap normally back to the left margin.")
    p2.paragraph_format.first_line_indent = Pt(36)
    doc.save(path)


def classic16_line_spacing(path):
    """Different line spacing options."""
    doc = Document()
    doc.add_heading("Line Spacing Test", level=2)
    text = "The quick brown fox jumps over the lazy dog. Pack my box with five dozen liquor jugs. How vexingly quick daft zebras jump."
    from docx.shared import Pt as PtShared
    for spacing_val, label in [(1.0, "Single"), (1.5, "1.5 Lines"), (2.0, "Double")]:
        doc.add_paragraph(f"{label} spacing:")
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing = spacing_val
        doc.add_paragraph()
    doc.save(path)


def classic17_page_break(path):
    """Document with explicit page breaks."""
    doc = Document()
    doc.add_heading("Page 1", level=1)
    doc.add_paragraph("Content on the first page.")
    doc.add_page_break()
    doc.add_heading("Page 2", level=1)
    doc.add_paragraph("Content on the second page after a page break.")
    doc.add_page_break()
    doc.add_heading("Page 3", level=1)
    doc.add_paragraph("Content on the third page.")
    doc.save(path)


def classic18_embedded_image(path):
    """Document with an embedded PNG image."""
    doc = Document()
    doc.add_heading("Image Test", level=2)
    doc.add_paragraph("Below is an embedded test image:")
    img_buf = _create_test_png(200, 100, (70, 130, 180))
    doc.add_picture(img_buf, width=Inches(3))
    doc.add_paragraph("The image above should be a blue rectangle.")
    doc.save(path)


def classic19_multiple_images(path):
    """Multiple images at different sizes."""
    doc = Document()
    doc.add_heading("Multiple Images", level=2)
    colors = [(220, 50, 50), (50, 180, 50), (50, 50, 220)]
    names = ["Red", "Green", "Blue"]
    for color, name in zip(colors, names):
        doc.add_paragraph(f"{name} image:")
        img_buf = _create_test_png(160, 80, color)
        doc.add_picture(img_buf, width=Inches(2.5))
    doc.save(path)


def classic20_table_with_many_rows(path):
    """Large table with 20 data rows."""
    doc = Document()
    doc.add_heading("Product Catalog", level=2)
    table = doc.add_table(rows=21, cols=4)
    table.style = "Table Grid"
    headers = ["ID", "Product", "Category", "Price"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    products = [
        ("Laptop", "Electronics", "$999"),
        ("Mouse", "Accessories", "$29"),
        ("Keyboard", "Accessories", "$59"),
        ("Monitor", "Electronics", "$349"),
        ("Headphones", "Audio", "$149"),
        ("Webcam", "Electronics", "$79"),
        ("USB Hub", "Accessories", "$25"),
        ("Desk Lamp", "Office", "$45"),
        ("Chair", "Furniture", "$299"),
        ("Standing Desk", "Furniture", "$599"),
        ("Printer", "Electronics", "$199"),
        ("Scanner", "Electronics", "$129"),
        ("Router", "Networking", "$89"),
        ("Cable Kit", "Accessories", "$19"),
        ("Mousepad", "Accessories", "$15"),
        ("Surge Protector", "Electronics", "$35"),
        ("External SSD", "Storage", "$109"),
        ("Flash Drive", "Storage", "$12"),
        ("Drawing Tablet", "Electronics", "$249"),
        ("Microphone", "Audio", "$179"),
    ]
    for ri, (prod, cat, price) in enumerate(products):
        table.rows[ri + 1].cells[0].text = str(ri + 1)
        table.rows[ri + 1].cells[1].text = prod
        table.rows[ri + 1].cells[2].text = cat
        table.rows[ri + 1].cells[3].text = price
    doc.save(path)


def classic21_nested_lists(path):
    """Nested bullet list (simulated with indentation)."""
    doc = Document()
    doc.add_heading("Project Structure", level=2)
    items = [
        (0, "src/"),
        (1, "MiniPdf/"),
        (2, "MiniPdf.cs"),
        (2, "PdfDocument.cs"),
        (2, "PdfWriter.cs"),
        (1, "MiniPdf.Tests/"),
        (2, "DocxToPdfConverterTests.cs"),
        (0, "scripts/"),
        (1, "Run-Benchmark.ps1"),
        (0, "README.md"),
    ]
    for level, text in items:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Pt(18 * level)
    doc.save(path)


def classic22_horizontal_rule(path):
    """Document with horizontal line separators (simulated)."""
    doc = Document()
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Content for section A goes here with enough text to see the layout.")
    _add_horizontal_rule(doc)
    doc.add_heading("Section B", level=2)
    doc.add_paragraph("Content for section B goes here below the horizontal divider.")
    _add_horizontal_rule(doc)
    doc.add_heading("Section C", level=2)
    doc.add_paragraph("Final section content.")
    doc.save(path)


def classic23_mixed_formatting_runs(path):
    """Single paragraph with many formatting runs."""
    doc = Document()
    doc.add_heading("Mixed Formatting", level=2)
    p = doc.add_paragraph()
    p.add_run("Normal, ")
    r1 = p.add_run("BOLD, ")
    r1.bold = True
    r2 = p.add_run("italic, ")
    r2.italic = True
    r3 = p.add_run("large, ")
    r3.font.size = Pt(18)
    r4 = p.add_run("small, ")
    r4.font.size = Pt(8)
    r5 = p.add_run("RED, ")
    r5.font.color.rgb = RGBColor(255, 0, 0)
    r6 = p.add_run("underlined.")
    r6.underline = True
    doc.save(path)


def classic24_two_column_table_layout(path):
    """Two-column layout using a borderless table."""
    doc = Document()
    doc.add_heading("Two-Column Layout", level=2)
    table = doc.add_table(rows=1, cols=2)
    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.text = (
        "Left column content. This is the first column of a two-column layout. "
        "It demonstrates how tables can be used for text layout purposes."
    )
    right_cell.text = (
        "Right column content. This is the second column. "
        "Both columns should render side-by-side in the PDF output."
    )
    doc.save(path)


def classic25_title_and_subtitle(path):
    """Title page with title and subtitle styles."""
    doc = Document()
    doc.add_paragraph("MiniPdf Benchmark Report", style="Title")
    doc.add_paragraph("Automated DOCX-to-PDF Conversion Quality Assessment", style="Subtitle")
    doc.add_paragraph()
    doc.add_paragraph("Prepared by: MiniPdf Team")
    doc.add_paragraph("Date: March 2026")
    doc.add_page_break()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This document tests the Title and Subtitle styles in MiniPdf conversion.")
    doc.save(path)


def classic26_table_alignment(path):
    """Table with aligned cell content (left, center, right)."""
    doc = Document()
    doc.add_heading("Cell Alignment Test", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = [("Left", WD_ALIGN_PARAGRAPH.LEFT),
               ("Center", WD_ALIGN_PARAGRAPH.CENTER),
               ("Right", WD_ALIGN_PARAGRAPH.RIGHT)]
    for ci, (h, align) in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].alignment = align
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ["Alice", "Engineering", "$95,000"],
        ["Bob", "Marketing", "$82,000"],
        ["Carol", "Finance", "$88,000"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            cell.paragraphs[0].alignment = headers[ci][1]
    doc.save(path)


def classic27_long_paragraph_wrapping(path):
    """Very long paragraph to test word wrapping."""
    doc = Document()
    doc.add_heading("Word Wrapping Test", level=2)
    long_text = (
        "This is a very long paragraph designed to test how MiniPdf handles word wrapping across "
        "line boundaries. The text should flow naturally from one line to the next without any "
        "awkward breaks or overflow. " * 10
    )
    doc.add_paragraph(long_text)
    doc.save(path)


def classic28_special_characters(path):
    """Document with special characters and symbols."""
    doc = Document()
    doc.add_heading("Special Characters", level=2)
    doc.add_paragraph("Ampersand: &, Less-than: <, Greater-than: >, Quotes: \" '")
    doc.add_paragraph("Copyright: \u00a9, Registered: \u00ae, Trademark: \u2122")
    doc.add_paragraph("Em-dash: \u2014, En-dash: \u2013, Ellipsis: \u2026")
    doc.add_paragraph("Currency: $ \u20ac \u00a3 \u00a5")
    doc.add_paragraph("Math: \u00b1 \u00d7 \u00f7 \u2260 \u2264 \u2265 \u221e")
    doc.save(path)


def classic29_table_with_image(path):
    """Table with an image in one cell."""
    doc = Document()
    doc.add_heading("Product Card", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Product"
    table.rows[0].cells[1].text = "Description"
    # Add image to first column of second row
    img_buf = _create_test_png(100, 60, (34, 139, 34))
    table.rows[1].cells[0].paragraphs[0].add_run().add_picture(img_buf, width=Inches(1.5))
    table.rows[1].cells[1].text = "MiniPdf Widget - A compact, efficient tool for PDF conversion. Lightweight and dependency-free."
    doc.save(path)


def classic30_comprehensive_report(path):
    """A comprehensive document combining many features."""
    doc = Document()
    # Title
    doc.add_paragraph("Annual Technology Report 2026", style="Title")
    doc.add_paragraph("A Comprehensive Overview", style="Subtitle")
    doc.add_page_break()

    # Table of contents placeholder
    doc.add_heading("Table of Contents", level=1)
    for i, title in enumerate(["Executive Summary", "Market Analysis", "Technology Trends", "Financial Overview", "Recommendations"], 1):
        doc.add_paragraph(f"{i}. {title}")

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report provides a comprehensive analysis of the technology landscape in 2026. "
        "Key findings include continued growth in AI adoption, increased focus on sustainability, "
        "and emerging trends in quantum computing."
    )

    # Market Analysis with table
    doc.add_heading("2. Market Analysis", level=1)
    doc.add_paragraph("The following table summarizes market share across key sectors:")
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Sector", "Market Share", "Growth"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    sectors = [
        ("Cloud Computing", "34%", "+12%"),
        ("AI/ML", "28%", "+23%"),
        ("Cybersecurity", "22%", "+18%"),
        ("IoT", "16%", "+8%"),
    ]
    for ri, (sector, share, growth) in enumerate(sectors):
        table.rows[ri + 1].cells[0].text = sector
        table.rows[ri + 1].cells[1].text = share
        table.rows[ri + 1].cells[2].text = growth

    # Technology Trends
    doc.add_heading("3. Technology Trends", level=1)
    doc.add_paragraph("Key trends identified:")
    for trend in ["Generative AI integration in enterprise software",
                  "Edge computing for real-time processing",
                  "Green technology and sustainable computing",
                  "Zero-trust security architectures",
                  "Low-code/no-code platform expansion"]:
        doc.add_paragraph(trend, style="List Bullet")

    # Image section
    doc.add_heading("4. Visual Summary", level=2)
    doc.add_paragraph("Growth indicator chart (placeholder):")
    img_buf = _create_test_png(300, 150, (46, 84, 150))
    doc.add_picture(img_buf, width=Inches(4))

    # Recommendations
    doc.add_heading("5. Recommendations", level=1)
    for i, rec in enumerate(["Invest in AI-driven automation tools",
                              "Prioritize cloud-native architectures",
                              "Strengthen cybersecurity posture",
                              "Explore quantum computing partnerships"], 1):
        doc.add_paragraph(rec, style="List Number")

    doc.save(path)


# ── Classic docx generators (31–60) ─────────────────────────────────────


def classic31_strikethrough_text(path):
    """Text with strikethrough and double-strikethrough formatting."""
    doc = Document()
    doc.add_heading("Strikethrough Test", level=2)
    p = doc.add_paragraph()
    p.add_run("Normal text. ")
    r1 = p.add_run("Single strikethrough. ")
    r1.font.strike = True
    r2 = p.add_run("Double strikethrough. ")
    r2.font.double_strike = True
    p.add_run("Back to normal.")

    doc.add_paragraph()
    p2 = doc.add_paragraph("Task list:")
    tasks = [
        ("Buy groceries", True),
        ("Write report", True),
        ("Schedule meeting", False),
        ("Review PR", False),
        ("Deploy to production", True),
    ]
    for task, done in tasks:
        tp = doc.add_paragraph()
        r = tp.add_run(task)
        if done:
            r.font.strike = True
    doc.save(path)


def classic32_superscript_subscript(path):
    """Superscript and subscript text."""
    doc = Document()
    doc.add_heading("Superscript and Subscript", level=2)

    p1 = doc.add_paragraph()
    p1.add_run("Einstein's equation: E = mc")
    r_sup = p1.add_run("2")
    r_sup.font.superscript = True

    p2 = doc.add_paragraph()
    p2.add_run("Water: H")
    r_sub1 = p2.add_run("2")
    r_sub1.font.subscript = True
    p2.add_run("O")

    p3 = doc.add_paragraph()
    p3.add_run("Carbon dioxide: CO")
    r_sub2 = p3.add_run("2")
    r_sub2.font.subscript = True

    p4 = doc.add_paragraph()
    p4.add_run("Footnote reference")
    r_fn = p4.add_run("1")
    r_fn.font.superscript = True

    p5 = doc.add_paragraph()
    p5.add_run("x")
    r_n = p5.add_run("n")
    r_n.font.superscript = True
    p5.add_run(" + y")
    r_n2 = p5.add_run("n")
    r_n2.font.superscript = True
    p5.add_run(" = z")
    r_n3 = p5.add_run("n")
    r_n3.font.superscript = True
    doc.save(path)


def classic33_highlighted_text(path):
    """Text with highlight colors."""
    doc = Document()
    doc.add_heading("Highlighted Text", level=2)
    from docx.enum.text import WD_COLOR_INDEX
    highlights = [
        ("Yellow highlight", WD_COLOR_INDEX.YELLOW),
        ("Green highlight", WD_COLOR_INDEX.GREEN),
        ("Turquoise highlight", WD_COLOR_INDEX.TURQUOISE),
        ("Pink highlight", WD_COLOR_INDEX.PINK),
        ("Red highlight", WD_COLOR_INDEX.RED),
        ("Blue highlight", WD_COLOR_INDEX.BLUE),
    ]
    for text, hl_color in highlights:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.highlight_color = hl_color
    doc.save(path)


def classic34_paragraph_borders(path):
    """Paragraphs with border boxes."""
    doc = Document()
    doc.add_heading("Bordered Paragraphs", level=2)

    # Create a paragraph with a box border
    for label, color in [("Important Notice", "FF0000"), ("Info Box", "4472C4"), ("Success", "00B050")]:
        p = doc.add_paragraph(label + ": This paragraph has a colored border around it.")
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        for side in ["top", "left", "bottom", "right"]:
            bdr = pBdr.makeelement(qn(f"w:{side}"), {
                qn("w:val"): "single",
                qn("w:sz"): "8",
                qn("w:space"): "4",
                qn("w:color"): color,
            })
            pBdr.append(bdr)
        pPr.append(pBdr)
        doc.add_paragraph()
    doc.save(path)


def classic35_tab_stops(path):
    """Paragraphs with custom tab stops."""
    doc = Document()
    doc.add_heading("Tab Stop Alignment", level=2)

    # Use tab characters with right-aligned tab stops to simulate a TOC
    entries = [
        ("Chapter 1: Introduction", "1"),
        ("Chapter 2: Getting Started", "5"),
        ("Chapter 3: Advanced Topics", "15"),
        ("Chapter 4: Best Practices", "28"),
        ("Chapter 5: Conclusion", "35"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.add_run(title)
        p.add_run("\t")
        p.add_run(page)
        # Add a right-aligned tab stop
        pPr = p._element.get_or_add_pPr()
        tabs = pPr.makeelement(qn("w:tabs"), {})
        tab = tabs.makeelement(qn("w:tab"), {
            qn("w:val"): "right",
            qn("w:leader"): "dot",
            qn("w:pos"): "9360",  # ~6.5 inches in twips
        })
        tabs.append(tab)
        pPr.append(tabs)
    doc.save(path)


def classic36_wide_table(path):
    """Table with many columns (8 columns)."""
    doc = Document()
    doc.add_heading("Weekly Schedule", level=2)
    table = doc.add_table(rows=6, cols=8)
    table.style = "Table Grid"
    headers = ["Time", "Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    schedule = [
        ["9:00", "Math", "English", "Science", "Math", "Art", "Free", "Free"],
        ["10:00", "English", "Math", "English", "Science", "Music", "Sports", "Free"],
        ["11:00", "Science", "Science", "Math", "English", "PE", "Free", "Free"],
        ["13:00", "History", "Art", "History", "Music", "Lab", "Free", "Free"],
        ["14:00", "PE", "Music", "Art", "History", "Free", "Free", "Free"],
    ]
    for ri, row_data in enumerate(schedule):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = val
    doc.save(path)


def classic37_nested_table(path):
    """Table inside a table cell."""
    doc = Document()
    doc.add_heading("Nested Table Layout", level=2)
    outer = doc.add_table(rows=2, cols=2)
    outer.style = "Table Grid"
    outer.rows[0].cells[0].text = "Section A"
    outer.rows[0].cells[1].text = "Section B"

    # Nested table in bottom-left cell
    cell = outer.rows[1].cells[0]
    cell.paragraphs[0].text = "Details:"
    inner = cell.add_table(rows=3, cols=2)
    inner.style = "Table Grid"
    inner.rows[0].cells[0].text = "Item"
    inner.rows[0].cells[1].text = "Qty"
    inner.rows[1].cells[0].text = "Widget"
    inner.rows[1].cells[1].text = "10"
    inner.rows[2].cells[0].text = "Gadget"
    inner.rows[2].cells[1].text = "5"

    outer.rows[1].cells[1].text = "This cell contains plain text while the adjacent cell has a nested table."
    doc.save(path)


def classic38_table_column_widths(path):
    """Table with explicit column widths."""
    doc = Document()
    doc.add_heading("Custom Column Widths", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    widths = [Inches(0.5), Inches(3.0), Inches(1.5), Inches(1.5)]
    for row in table.rows:
        for ci, w in enumerate(widths):
            row.cells[ci].width = w
    headers = ["#", "Description", "Category", "Amount"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    data = [
        ["1", "Office supplies and stationery", "Operations", "$245.00"],
        ["2", "Cloud hosting monthly fee", "Technology", "$1,200.00"],
        ["3", "Team lunch and catering", "Meals", "$380.00"],
        ["4", "Conference registration", "Travel", "$599.00"],
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = val
    doc.save(path)


def classic39_financial_report(path):
    """Financial report with colored variance values."""
    doc = Document()
    doc.add_heading("Financial Summary Q4 2025", level=1)
    doc.add_paragraph("All amounts in USD thousands.")

    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    headers = ["Line Item", "Budget", "Actual", "Variance"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    data = [
        ("Revenue", "$500", "$520", "+$20"),
        ("COGS", "$200", "$210", "-$10"),
        ("Gross Profit", "$300", "$310", "+$10"),
        ("Operating Exp", "$150", "$140", "+$10"),
        ("Net Income", "$150", "$170", "+$20"),
        ("EPS", "$1.50", "$1.70", "+$0.20"),
    ]
    for ri, (item, budget, actual, variance) in enumerate(data):
        row = table.rows[ri + 1]
        row.cells[0].text = item
        row.cells[1].text = budget
        row.cells[2].text = actual
        p = row.cells[3].paragraphs[0]
        r = p.add_run(variance)
        if variance.startswith("+"):
            r.font.color.rgb = RGBColor(0, 128, 0)
        else:
            r.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph()
    doc.add_paragraph("Note: Positive variance indicates favorable performance.")
    doc.save(path)


def classic40_resume(path):
    """Simple resume/CV document."""
    doc = Document()
    doc.add_paragraph("JOHN DOE", style="Title")
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.add_run("john.doe@email.com | +1-555-0100 | New York, NY")

    _add_horizontal_rule(doc)

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Experienced software engineer with 8+ years of expertise in building scalable "
        "web applications and distributed systems. Proficient in C#, Python, and JavaScript."
    )

    doc.add_heading("Experience", level=2)
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Senior Software Engineer")
    r_title.bold = True
    p_title.add_run(" - Tech Corp Inc.")
    p_date = doc.add_paragraph()
    r_date = p_date.add_run("January 2020 - Present")
    r_date.italic = True
    for item in ["Led team of 5 engineers on microservices migration",
                 "Reduced API latency by 40% through caching optimization",
                 "Implemented CI/CD pipeline using GitHub Actions"]:
        doc.add_paragraph(item, style="List Bullet")

    p_title2 = doc.add_paragraph()
    r_title2 = p_title2.add_run("Software Engineer")
    r_title2.bold = True
    p_title2.add_run(" - StartupXYZ")
    p_date2 = doc.add_paragraph()
    r_date2 = p_date2.add_run("June 2016 - December 2019")
    r_date2.italic = True
    for item in ["Built RESTful APIs serving 1M+ daily requests",
                 "Developed real-time notification system using WebSockets"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Education", level=2)
    p_edu = doc.add_paragraph()
    r_edu = p_edu.add_run("B.S. Computer Science")
    r_edu.bold = True
    p_edu.add_run(" - State University, 2016")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph("C#, Python, JavaScript, TypeScript, SQL, Docker, Kubernetes, Azure, AWS")
    doc.save(path)


def classic41_business_letter(path):
    """Formal business letter layout."""
    doc = Document()
    # Sender info - right aligned
    for line in ["ACME Corporation", "123 Business Ave, Suite 500", "New York, NY 10001"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph("March 1, 2026")
    doc.add_paragraph()

    # Recipient
    for line in ["Mr. James Wilson", "Widget Industries", "456 Commerce St", "San Francisco, CA 94102"]:
        doc.add_paragraph(line)
    doc.add_paragraph()

    doc.add_paragraph("Dear Mr. Wilson,")
    doc.add_paragraph()

    doc.add_paragraph(
        "Thank you for your interest in our products. We are pleased to inform you that "
        "your order #ORD-2026-0315 has been processed and is scheduled for delivery by "
        "March 15, 2026."
    )
    doc.add_paragraph(
        "Please find enclosed the detailed invoice and shipping confirmation. If you have "
        "any questions regarding your order, please do not hesitate to contact our customer "
        "service team at support@acme.com or call us at +1-555-0200."
    )
    doc.add_paragraph(
        "We value your business and look forward to a continued partnership."
    )

    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    r_sig = p_sig.add_run("Sarah Johnson")
    r_sig.bold = True
    doc.add_paragraph("Vice President of Sales")
    doc.add_paragraph("ACME Corporation")
    doc.save(path)


def classic42_meeting_minutes(path):
    """Meeting minutes document."""
    doc = Document()
    doc.add_heading("Meeting Minutes", level=1)
    doc.add_paragraph()

    # Meeting details table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    details = [
        ("Date", "March 3, 2026"),
        ("Time", "10:00 AM - 11:30 AM"),
        ("Location", "Conference Room B"),
        ("Attendees", "Alice, Bob, Carol, David, Eve"),
    ]
    for ri, (label, value) in enumerate(details):
        cell_l = table.rows[ri].cells[0]
        cell_l.text = label
        for run in cell_l.paragraphs[0].runs:
            run.bold = True
        _set_cell_shading(cell_l, "E2EFDA")
        table.rows[ri].cells[1].text = value

    doc.add_paragraph()
    doc.add_heading("Agenda Items", level=2)
    for i, item in enumerate(["Q4 Review", "Budget Planning", "New Hire Onboarding", "Action Items"], 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_heading("Discussion Summary", level=2)
    doc.add_paragraph(
        "Alice presented the Q4 results showing a 15% revenue increase. "
        "Bob proposed reallocating 10% of the marketing budget to R&D. "
        "Carol reported that three new engineering positions have been approved."
    )

    doc.add_heading("Action Items", level=2)
    actions = [
        ("Bob", "Submit revised budget proposal", "March 10"),
        ("Carol", "Post job listings for engineering roles", "March 7"),
        ("David", "Prepare onboarding materials", "March 14"),
        ("Eve", "Schedule follow-up meeting", "March 5"),
    ]
    t2 = doc.add_table(rows=len(actions) + 1, cols=3)
    t2.style = "Table Grid"
    for i, h in enumerate(["Owner", "Action", "Due Date"]):
        cell = t2.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ri, (owner, action, due) in enumerate(actions):
        t2.rows[ri + 1].cells[0].text = owner
        t2.rows[ri + 1].cells[1].text = action
        t2.rows[ri + 1].cells[2].text = due
    doc.save(path)


def classic43_invoice_document(path):
    """Invoice document with line items and totals."""
    doc = Document()
    # Company header
    p_company = doc.add_paragraph()
    r_company = p_company.add_run("ACME SOLUTIONS INC.")
    r_company.bold = True
    r_company.font.size = Pt(18)
    doc.add_paragraph("789 Tech Boulevard, Austin, TX 78701")
    doc.add_paragraph("Phone: +1-555-0300 | Email: billing@acme-solutions.com")

    _add_horizontal_rule(doc)

    p_inv = doc.add_paragraph()
    r_inv = p_inv.add_run("INVOICE")
    r_inv.bold = True
    r_inv.font.size = Pt(14)
    p_inv.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    # Invoice details
    details_table = doc.add_table(rows=3, cols=4)
    details_table.rows[0].cells[0].text = "Invoice #:"
    details_table.rows[0].cells[1].text = "INV-2026-0087"
    details_table.rows[0].cells[2].text = "Date:"
    details_table.rows[0].cells[3].text = "March 1, 2026"
    details_table.rows[1].cells[0].text = "Due Date:"
    details_table.rows[1].cells[1].text = "March 31, 2026"
    details_table.rows[1].cells[2].text = "Terms:"
    details_table.rows[1].cells[3].text = "Net 30"
    details_table.rows[2].cells[0].text = "Bill To:"
    details_table.rows[2].cells[1].text = "Widget Industries"
    details_table.rows[2].cells[2].text = "Ship To:"
    details_table.rows[2].cells[3].text = "Same as billing"

    doc.add_paragraph()

    # Line items
    items_table = doc.add_table(rows=6, cols=5)
    items_table.style = "Table Grid"
    headers = ["Item", "Description", "Qty", "Unit Price", "Total"]
    for i, h in enumerate(headers):
        cell = items_table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    line_items = [
        ("SVC-001", "Consulting Services (40 hrs)", "40", "$150.00", "$6,000.00"),
        ("LIC-002", "Enterprise License (Annual)", "5", "$499.00", "$2,495.00"),
        ("HW-003", "Server Hardware", "2", "$2,499.00", "$4,998.00"),
        ("SUP-004", "Premium Support Plan", "1", "$1,800.00", "$1,800.00"),
        ("TRN-005", "On-site Training (2 days)", "1", "$3,000.00", "$3,000.00"),
    ]
    for ri, items in enumerate(line_items):
        for ci, val in enumerate(items):
            items_table.rows[ri + 1].cells[ci].text = val

    doc.add_paragraph()
    # Totals - right aligned
    for label, amount in [("Subtotal:", "$18,293.00"), ("Tax (8.25%):", "$1,509.17"), ("Total Due:", "$19,802.17")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"{label}  ")
        r_amt = p.add_run(amount)
        if label.startswith("Total"):
            r_amt.bold = True
            r_amt.font.size = Pt(14)
    doc.save(path)


def classic44_memo(path):
    """Internal memo document."""
    doc = Document()
    p_memo = doc.add_paragraph()
    r_memo = p_memo.add_run("MEMORANDUM")
    r_memo.bold = True
    r_memo.font.size = Pt(16)
    p_memo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_horizontal_rule(doc)

    fields = [
        ("TO:", "All Department Heads"),
        ("FROM:", "Maria Garcia, CEO"),
        ("DATE:", "March 3, 2026"),
        ("RE:", "Quarterly Performance Review Process Changes")
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        r_label = p.add_run(label + "  ")
        r_label.bold = True
        p.add_run(value)

    _add_horizontal_rule(doc)

    doc.add_paragraph(
        "Effective immediately, we are implementing several changes to our quarterly "
        "performance review process. These changes are designed to streamline evaluation "
        "procedures and provide more actionable feedback to team members."
    )
    doc.add_paragraph()
    doc.add_heading("Key Changes", level=2)
    changes = [
        "Reviews will now be conducted bi-monthly instead of quarterly",
        "Self-assessment forms must be submitted 5 business days before the review",
        "360-degree feedback will be incorporated for all managerial positions",
        "New rating scale: 1-5 (replacing the current A-F system)",
        "All reviews must be completed within a 2-week window",
    ]
    for change in changes:
        doc.add_paragraph(change, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Please share this information with your teams and direct any questions "
        "to the HR department at hr@company.com."
    )
    doc.save(path)


def classic45_project_plan(path):
    """Project plan with timeline table."""
    doc = Document()
    doc.add_heading("Project Plan: Website Redesign", level=1)
    doc.add_paragraph("Project Manager: Sarah Chen | Start Date: March 2026")

    doc.add_heading("Project Overview", level=2)
    doc.add_paragraph(
        "This project aims to redesign the company website to improve user experience, "
        "modernize the visual design, and optimize for mobile devices."
    )

    doc.add_heading("Timeline", level=2)
    table = doc.add_table(rows=8, cols=5)
    table.style = "Table Grid"
    headers = ["Phase", "Task", "Owner", "Start", "End"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "4472C4")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    tasks = [
        ("Discovery", "User research & interviews", "UX Team", "Mar 1", "Mar 14"),
        ("Discovery", "Competitive analysis", "Marketing", "Mar 1", "Mar 7"),
        ("Design", "Wireframes", "Design Team", "Mar 15", "Mar 28"),
        ("Design", "Visual mockups", "Design Team", "Mar 29", "Apr 11"),
        ("Development", "Frontend build", "Dev Team", "Apr 12", "May 9"),
        ("Development", "Backend integration", "Dev Team", "Apr 19", "May 16"),
        ("Launch", "QA testing & deployment", "QA Team", "May 17", "May 30"),
    ]
    for ri, (phase, task, owner, start, end) in enumerate(tasks):
        row = table.rows[ri + 1]
        row.cells[0].text = phase
        row.cells[1].text = task
        row.cells[2].text = owner
        row.cells[3].text = start
        row.cells[4].text = end

    doc.add_heading("Budget", level=2)
    budget_table = doc.add_table(rows=5, cols=2)
    budget_table.style = "Table Grid"
    budget_table.rows[0].cells[0].text = "Category"
    budget_table.rows[0].cells[1].text = "Amount"
    for run in budget_table.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
    for run in budget_table.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True
    budget = [("Design", "$15,000"), ("Development", "$45,000"), ("QA & Testing", "$8,000"), ("Total", "$68,000")]
    for ri, (cat, amt) in enumerate(budget):
        budget_table.rows[ri + 1].cells[0].text = cat
        budget_table.rows[ri + 1].cells[1].text = amt
        if cat == "Total":
            for cell in budget_table.rows[ri + 1].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.save(path)


def classic46_comparison_table(path):
    """Product comparison table with checkmarks and crosses."""
    doc = Document()
    doc.add_heading("Product Comparison", level=1)

    table = doc.add_table(rows=9, cols=4)
    table.style = "Table Grid"
    headers = ["Feature", "Basic", "Pro", "Enterprise"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    features = [
        ("Cloud Storage", "5 GB", "50 GB", "Unlimited"),
        ("Users", "1", "10", "Unlimited"),
        ("API Access", "No", "Yes", "Yes"),
        ("Priority Support", "No", "Yes", "Yes"),
        ("Custom Branding", "No", "No", "Yes"),
        ("SLA Guarantee", "No", "99.9%", "99.99%"),
        ("Data Export", "CSV", "CSV, JSON", "All formats"),
        ("Price/month", "$9", "$29", "$99"),
    ]
    for ri, (feature, basic, pro, ent) in enumerate(features):
        row = table.rows[ri + 1]
        row.cells[0].text = feature
        for ci, val in enumerate([basic, pro, ent], 1):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(path)


def classic47_data_dictionary(path):
    """Database data dictionary document."""
    doc = Document()
    doc.add_heading("Data Dictionary", level=1)
    doc.add_paragraph("Database: CustomerDB | Version: 2.1 | Last Updated: March 2026")

    tables_spec = [
        ("customers", [
            ("id", "INT", "PRIMARY KEY", "Unique customer identifier"),
            ("name", "VARCHAR(100)", "NOT NULL", "Full name"),
            ("email", "VARCHAR(255)", "UNIQUE", "Email address"),
            ("created_at", "DATETIME", "DEFAULT NOW()", "Account creation date"),
            ("status", "ENUM", "DEFAULT 'active'", "Account status"),
        ]),
        ("orders", [
            ("id", "INT", "PRIMARY KEY", "Order identifier"),
            ("customer_id", "INT", "FOREIGN KEY", "Reference to customers.id"),
            ("total", "DECIMAL(10,2)", "NOT NULL", "Order total amount"),
            ("status", "VARCHAR(20)", "DEFAULT 'pending'", "Order status"),
            ("created_at", "DATETIME", "DEFAULT NOW()", "Order creation date"),
        ]),
    ]

    for table_name, columns in tables_spec:
        doc.add_heading(f"Table: {table_name}", level=2)
        t = doc.add_table(rows=len(columns) + 1, cols=4)
        t.style = "Table Grid"
        for i, h in enumerate(["Column", "Type", "Constraints", "Description"]):
            cell = t.rows[0].cells[i]
            cell.text = h
            _set_cell_shading(cell, "4472C4")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
        for ri, (col, dtype, constraint, desc) in enumerate(columns):
            row = t.rows[ri + 1]
            row.cells[0].text = col
            row.cells[1].text = dtype
            row.cells[2].text = constraint
            row.cells[3].text = desc
        doc.add_paragraph()
    doc.save(path)


def classic48_multi_level_headings(path):
    """Document with 4 levels of heading depth and content."""
    doc = Document()
    doc.add_heading("Software Architecture Document", level=1)

    for i in range(1, 4):
        doc.add_heading(f"{i}. Module {i}", level=1)
        doc.add_paragraph(f"Module {i} provides core functionality for the system.")

        for j in range(1, 3):
            doc.add_heading(f"{i}.{j} Component {j}", level=2)
            doc.add_paragraph(f"Component {i}.{j} handles specific operations within Module {i}.")

            for k in range(1, 3):
                doc.add_heading(f"{i}.{j}.{k} Sub-component", level=3)
                doc.add_paragraph(f"Detailed description of sub-component {i}.{j}.{k}.")

                doc.add_heading(f"{i}.{j}.{k}.1 Implementation Notes", level=4)
                doc.add_paragraph("Implementation uses factory pattern with dependency injection.")
    doc.save(path)


def classic49_cjk_document(path):
    """Document with Chinese, Japanese, and Korean text."""
    doc = Document()
    doc.add_heading("CJK Text Sample", level=1)

    doc.add_heading("Chinese (Simplified)", level=2)
    doc.add_paragraph("MiniPdf 是一个轻量级的 .NET 库，用于将 Office 文档转换为 PDF 格式。")
    doc.add_paragraph("它不依赖 Microsoft Office，可以在任何平台上运行。")

    doc.add_heading("Chinese (Traditional)", level=2)
    doc.add_paragraph("MiniPdf 是一個輕量級的 .NET 庫，用於將 Office 文檔轉換為 PDF 格式。")

    doc.add_heading("Japanese", level=2)
    doc.add_paragraph("MiniPdf は軽量な .NET ライブラリで、Office ドキュメントを PDF に変換します。")

    doc.add_heading("Korean", level=2)
    doc.add_paragraph("MiniPdf는 Office 문서를 PDF로 변환하는 경량 .NET 라이브러리입니다.")

    doc.add_heading("Mixed CJK Table", level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Language"
    table.rows[0].cells[1].text = "Hello"
    table.rows[0].cells[2].text = "Thank You"
    data = [
        ("Chinese", "你好", "谢谢"),
        ("Japanese", "こんにちは", "ありがとう"),
        ("Korean", "안녕하세요", "감사합니다"),
        ("English", "Hello", "Thank you"),
    ]
    for ri, (lang, hello, thanks) in enumerate(data):
        table.rows[ri + 1].cells[0].text = lang
        table.rows[ri + 1].cells[1].text = hello
        table.rows[ri + 1].cells[2].text = thanks
    doc.save(path)


def classic50_long_table_with_formatting(path):
    """Large table with 30 rows and alternating row colors."""
    doc = Document()
    doc.add_heading("Server Inventory Report", level=1)
    table = doc.add_table(rows=31, cols=5)
    table.style = "Table Grid"
    headers = ["Server ID", "Hostname", "IP Address", "OS", "Status"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    oses = ["Ubuntu 22.04", "Windows Server 2022", "RHEL 9", "Debian 12", "CentOS Stream 9"]
    statuses = ["Online", "Online", "Online", "Maintenance", "Online",
                "Online", "Offline", "Online", "Online", "Online"]
    for ri in range(30):
        row = table.rows[ri + 1]
        row.cells[0].text = f"SRV-{ri + 1:03d}"
        row.cells[1].text = f"server-{ri + 1:03d}.local"
        row.cells[2].text = f"10.0.{ri // 256}.{ri % 256 + 1}"
        row.cells[3].text = oses[ri % len(oses)]
        status = statuses[ri % len(statuses)]
        p = row.cells[4].paragraphs[0]
        r = p.add_run(status)
        if status == "Online":
            r.font.color.rgb = RGBColor(0, 128, 0)
        elif status == "Offline":
            r.font.color.rgb = RGBColor(255, 0, 0)
        else:
            r.font.color.rgb = RGBColor(255, 165, 0)
        if ri % 2 == 0:
            for ci in range(5):
                _set_cell_shading(row.cells[ci], "D9E2F3")
    doc.save(path)


def classic51_underline_styles(path):
    """Various underline styles."""
    doc = Document()
    doc.add_heading("Underline Styles", level=2)
    from docx.enum.text import WD_UNDERLINE
    styles = [
        ("Single underline", WD_UNDERLINE.SINGLE),
        ("Double underline", WD_UNDERLINE.DOUBLE),
        ("Thick underline", WD_UNDERLINE.THICK),
        ("Dotted underline", WD_UNDERLINE.DOTTED),
        ("Dash underline", WD_UNDERLINE.DASH),
        ("Wavy underline", WD_UNDERLINE.WAVY),
    ]
    for text, style in styles:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.underline = style
    doc.save(path)


def classic52_spacing_before_after(path):
    """Paragraphs with explicit before/after spacing."""
    doc = Document()
    doc.add_heading("Paragraph Spacing Test", level=2)

    spacings = [(0, 0), (6, 6), (12, 12), (18, 6), (6, 18), (24, 24)]
    for before, after in spacings:
        p = doc.add_paragraph(f"SpaceBefore={before}pt, SpaceAfter={after}pt. "
                              "Sample text to show paragraph spacing effects.")
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
    doc.save(path)


def classic53_table_merged_complex(path):
    """Complex table with merged cells in rows and columns."""
    doc = Document()
    doc.add_heading("Course Schedule", level=1)

    table = doc.add_table(rows=6, cols=5)
    table.style = "Table Grid"
    headers = ["", "Period 1", "Period 2", "Period 3", "Period 4"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "4472C4")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    for ri, day in enumerate(days):
        table.rows[ri + 1].cells[0].text = day
        _set_cell_shading(table.rows[ri + 1].cells[0], "D9E2F3")

    # Fill in schedule
    schedule_data = [
        ["Math", "English", "Science", "Art"],
        ["English", "Math", "History", "PE"],
        ["Science", "Science", "Math", "Music"],
        ["History", "Art", "English", "Lab"],
        ["PE", "Music", "Math", "English"],
    ]
    for ri, row_data in enumerate(schedule_data):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci + 1].text = val

    # Merge Wednesday periods 2-3 for "Science Lab"
    merged = table.rows[3].cells[1].merge(table.rows[3].cells[2])
    merged.text = "Science Lab (Double Period)"
    doc.save(path)


def classic54_multi_font_family(path):
    """Text with different font families."""
    doc = Document()
    doc.add_heading("Font Family Showcase", level=2)
    fonts = [
        ("Arial", "The quick brown fox jumps over the lazy dog."),
        ("Times New Roman", "The quick brown fox jumps over the lazy dog."),
        ("Courier New", "The quick brown fox jumps over the lazy dog."),
        ("Calibri", "The quick brown fox jumps over the lazy dog."),
        ("Verdana", "The quick brown fox jumps over the lazy dog."),
        ("Georgia", "The quick brown fox jumps over the lazy dog."),
    ]
    for font_name, text in fonts:
        p = doc.add_paragraph()
        r = p.add_run(f"{font_name}: {text}")
        r.font.name = font_name
    doc.save(path)


def classic55_background_shading_paragraph(path):
    """Paragraphs with background shading colors."""
    doc = Document()
    doc.add_heading("Paragraph Shading", level=2)

    shadings = [
        ("Light blue background", "D9E2F3"),
        ("Light green background", "E2EFDA"),
        ("Light yellow background", "FFF2CC"),
        ("Light red background", "FCE4EC"),
        ("Light gray background", "F2F2F2"),
    ]
    for text, color in shadings:
        p = doc.add_paragraph(text + " - This paragraph has a colored background.")
        # Apply shading via XML
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color,
        })
        pPr.append(shd)
    doc.save(path)


def classic56_images_and_tables_mixed(path):
    """Document with alternating images and tables."""
    doc = Document()
    doc.add_heading("Product Catalog", level=1)

    products = [
        ("Widget Pro", (50, 120, 200), "$49.99", "Premium widget with advanced features"),
        ("Gadget Max", (200, 80, 50), "$79.99", "Industrial-grade gadget for heavy use"),
        ("Connector Plus", (50, 180, 50), "$19.99", "Universal connector with fast transfer"),
    ]

    for name, color, price, desc in products:
        doc.add_heading(name, level=2)
        img_buf = _create_test_png(180, 90, color)
        doc.add_picture(img_buf, width=Inches(2.5))

        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Product"
        table.rows[0].cells[1].text = name
        table.rows[1].cells[0].text = "Price"
        table.rows[1].cells[1].text = price
        table.rows[2].cells[0].text = "Description"
        table.rows[2].cells[1].text = desc
        doc.add_paragraph()
    doc.save(path)


def classic57_right_to_left_text(path):
    """Document with right-to-left text simulation."""
    doc = Document()
    doc.add_heading("Text Direction Test", level=2)

    doc.add_paragraph("Left-to-right (default):")
    doc.add_paragraph("This is standard left-to-right English text.")

    doc.add_paragraph()
    doc.add_paragraph("Right-aligned text (simulating RTL):")
    p = doc.add_paragraph("This text is right-aligned to simulate right-to-left layout.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph("Hebrew sample (RTL):")
    p2 = doc.add_paragraph("\u05E9\u05DC\u05D5\u05DD \u05E2\u05D5\u05DC\u05DD")  # Shalom Olam
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph("Arabic sample (RTL):")
    p3 = doc.add_paragraph("\u0645\u0631\u062D\u0628\u0627 \u0628\u0627\u0644\u0639\u0627\u0644\u0645")  # Marhaba bilalam
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(path)


def classic58_dense_paragraph_document(path):
    """Long document with many paragraphs spanning multiple pages."""
    doc = Document()
    doc.add_heading("Research Paper: Modern Software Engineering", level=1)
    doc.add_paragraph("Author: Dr. Jane Smith | Published: March 2026")
    doc.add_paragraph()

    sections = [
        ("Abstract", 2),
        ("1. Introduction", 3),
        ("2. Literature Review", 4),
        ("3. Methodology", 3),
        ("4. Results", 4),
        ("5. Discussion", 3),
        ("6. Conclusion", 2),
    ]
    filler = (
        "Modern software engineering practices emphasize continuous integration, "
        "automated testing, and agile methodologies. The rapid evolution of cloud computing "
        "and containerization has transformed how teams build and deploy applications. "
        "Microservices architecture enables independent scaling and deployment of components. "
    )
    for title, para_count in sections:
        doc.add_heading(title, level=2)
        for _ in range(para_count):
            doc.add_paragraph(filler * 3)
    doc.save(path)


def classic59_numbered_and_bullet_mixed(path):
    """Mixed numbered and bullet lists."""
    doc = Document()
    doc.add_heading("Installation Guide", level=1)

    doc.add_heading("Prerequisites", level=2)
    for item in [".NET 8.0 SDK or later", "Visual Studio Code", "Git", "Python 3.10+"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Installation Steps", level=2)
    for step in ["Clone the repository from GitHub",
                 "Open the project in Visual Studio Code",
                 "Restore NuGet packages",
                 "Build the solution",
                 "Run the test suite"]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Configuration Options", level=2)
    for item in ["Set output directory in appsettings.json",
                 "Configure font embedding preferences",
                 "Enable or disable image compression"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Troubleshooting", level=2)
    for step in ["Verify .NET SDK installation with dotnet --version",
                 "Clear NuGet cache if packages fail to restore",
                 "Check file permissions on output directory"]:
        doc.add_paragraph(step, style="List Number")
    doc.save(path)


def classic60_comprehensive_styled_report(path):
    """Comprehensive report using many styling features together."""
    doc = Document()
    # Title page
    doc.add_paragraph("Technology Trends Report", style="Title")
    doc.add_paragraph("Q1 2026 Analysis", style="Subtitle")
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.add_run("Prepared by MiniPdf Analytics Team")
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("March 2026")
    r_date.italic = True
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    p_summary = doc.add_paragraph()
    p_summary.add_run("Key Findings: ").bold = True
    p_summary.add_run(
        "The technology sector continues to show strong growth driven by AI adoption, "
        "cloud migration, and digital transformation initiatives."
    )

    # Market Overview with styled table
    doc.add_heading("Market Overview", level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["Sector", "Q4 2025", "Q1 2026", "Change"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    sectors = [
        ("Cloud Computing", "$180B", "$195B", "+8.3%"),
        ("AI/ML", "$95B", "$112B", "+17.9%"),
        ("Cybersecurity", "$72B", "$78B", "+8.3%"),
        ("IoT", "$45B", "$48B", "+6.7%"),
        ("Blockchain", "$12B", "$11B", "-8.3%"),
    ]
    for ri, (sector, q4, q1, change) in enumerate(sectors):
        row = table.rows[ri + 1]
        row.cells[0].text = sector
        row.cells[1].text = q4
        row.cells[2].text = q1
        p = row.cells[3].paragraphs[0]
        r = p.add_run(change)
        if change.startswith("+"):
            r.font.color.rgb = RGBColor(0, 128, 0)
        else:
            r.font.color.rgb = RGBColor(255, 0, 0)
        if ri % 2 == 0:
            for ci in range(4):
                _set_cell_shading(row.cells[ci], "D9E2F3")

    # Trends
    doc.add_heading("Key Trends", level=1)
    doc.add_heading("Artificial Intelligence", level=2)
    for item in ["Large Language Models becoming mainstream",
                 "AI-assisted coding tools adoption growing",
                 "Regulation frameworks being established"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Cloud & Infrastructure", level=2)
    for item in ["Multi-cloud strategies dominating",
                 "Serverless architecture expanding",
                 "Edge computing gaining traction"]:
        doc.add_paragraph(item, style="List Bullet")

    # Visual section
    doc.add_heading("Visual Summary", level=1)
    doc.add_paragraph("Growth indicator (placeholder):")
    img_buf = _create_test_png(300, 120, (46, 84, 150))
    doc.add_picture(img_buf, width=Inches(4))

    # Recommendations with numbered list
    doc.add_heading("Recommendations", level=1)
    for rec in ["Accelerate AI integration strategy",
                "Diversify cloud vendor dependencies",
                "Invest in employee upskilling programs",
                "Strengthen data governance frameworks",
                "Evaluate emerging quantum computing capabilities"]:
        doc.add_paragraph(rec, style="List Number")

    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = p_footer.add_run("--- End of Report ---")
    r_footer.italic = True
    r_footer.font.color.rgb = RGBColor(128, 128, 128)
    doc.save(path)


# ── Classic docx generators (61–90) ─────────────────────────────────────


def classic61_header_and_footer(path):
    """Document with header and footer text."""
    doc = Document()
    doc.add_heading("Header and Footer Test", level=1)

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "MiniPdf Benchmark Report"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.runs[0]
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(128, 128, 128)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = "Page 1 | Confidential"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.runs[0]
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    for i in range(1, 6):
        doc.add_paragraph(
            f"Section {i}: This content appears between the header and footer. "
            "It demonstrates how MiniPdf handles page headers and footers in DOCX conversion."
        )
    doc.save(path)


def classic62_footnote_references(path):
    """Document with superscript footnote references and bottom notes."""
    doc = Document()
    doc.add_heading("Research with Footnotes", level=1)

    p1 = doc.add_paragraph()
    p1.add_run("The theory of relativity")
    r1 = p1.add_run("1")
    r1.font.superscript = True
    r1.font.size = Pt(8)
    p1.add_run(" fundamentally changed our understanding of space and time.")

    p2 = doc.add_paragraph()
    p2.add_run("Quantum mechanics")
    r2 = p2.add_run("2")
    r2.font.superscript = True
    r2.font.size = Pt(8)
    p2.add_run(" describes the behavior of particles at the atomic level.")

    p3 = doc.add_paragraph()
    p3.add_run("The Standard Model")
    r3 = p3.add_run("3")
    r3.font.superscript = True
    r3.font.size = Pt(8)
    p3.add_run(" classifies all known elementary particles.")

    doc.add_paragraph()
    _add_horizontal_rule(doc)
    notes = [
        "1. Einstein, A. (1905). On the Electrodynamics of Moving Bodies.",
        "2. Planck, M. (1900). On the Law of Distribution of Energy.",
        "3. Glashow, S. (1961). Partial-symmetries of Weak Interactions.",
    ]
    for note in notes:
        p = doc.add_paragraph(note)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(8)
    doc.save(path)


def classic63_toc_style_headings(path):
    """Document simulating a table of contents with hierarchical headings."""
    doc = Document()
    doc.add_paragraph("Table of Contents", style="Title")
    doc.add_paragraph()

    toc_entries = [
        (0, "Chapter 1: Introduction", "1"),
        (1, "1.1 Background", "3"),
        (1, "1.2 Objectives", "5"),
        (0, "Chapter 2: Literature Review", "7"),
        (1, "2.1 Historical Context", "8"),
        (1, "2.2 Current Research", "12"),
        (2, "2.2.1 Methodology", "13"),
        (2, "2.2.2 Findings", "15"),
        (0, "Chapter 3: Methodology", "18"),
        (1, "3.1 Data Collection", "19"),
        (1, "3.2 Analysis Framework", "22"),
        (0, "Chapter 4: Results", "25"),
        (0, "Chapter 5: Discussion", "30"),
        (0, "Chapter 6: Conclusion", "35"),
        (0, "References", "38"),
    ]

    table = doc.add_table(rows=len(toc_entries), cols=2)
    for ri, (level, title, page) in enumerate(toc_entries):
        indent = "    " * level
        cell_title = table.rows[ri].cells[0]
        cell_page = table.rows[ri].cells[1]
        p_title = cell_title.paragraphs[0]
        r_title = p_title.add_run(f"{indent}{title}")
        if level == 0:
            r_title.bold = True
        r_title.font.size = Pt(11 - level)
        p_page = cell_page.paragraphs[0]
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_page.add_run(page)
    doc.save(path)


def classic64_multi_column_layout(path):
    """Simulated two-column text layout using a table."""
    doc = Document()
    doc.add_heading("Newsletter - March 2026", level=1)

    table = doc.add_table(rows=1, cols=2)
    left_text = (
        "Artificial intelligence continues to reshape the technology landscape. "
        "New advances in large language models enable more natural human-computer "
        "interaction. Companies worldwide are investing heavily in AI research "
        "and development, anticipating transformative impacts across industries "
        "from healthcare to manufacturing."
    )
    right_text = (
        "Cloud computing has become the backbone of modern enterprise IT. "
        "Multi-cloud strategies allow organizations to leverage the best "
        "features of different providers while avoiding vendor lock-in. "
        "Edge computing supplements cloud by processing data closer to "
        "its source, reducing latency for critical applications."
    )
    table.rows[0].cells[0].text = left_text
    table.rows[0].cells[1].text = right_text

    doc.add_paragraph()
    doc.add_heading("Featured Article", level=2)
    doc.add_paragraph(
        "Open source software has become the foundation of modern software development. "
        "Projects like Linux, Kubernetes, and .NET have demonstrated how community-driven "
        "development can produce enterprise-grade software. The MiniPdf project itself is "
        "an example of this approach, providing PDF generation capabilities without proprietary "
        "dependencies."
    )
    doc.save(path)


def classic65_code_block_styling(path):
    """Document with styled code blocks using monospace fonts and shading."""
    doc = Document()
    doc.add_heading("Code Examples", level=1)

    doc.add_heading("C# Example", level=2)
    code_cs = (
        'using MiniPdf;\n\n'
        'var doc = new PdfDocument();\n'
        'var page = doc.AddPage();\n'
        'page.DrawText("Hello, PDF!", 72, 700);\n'
        'doc.Save("output.pdf");'
    )
    p = doc.add_paragraph()
    r = p.add_run(code_cs)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F5F5F5",
    })
    pPr.append(shd)

    doc.add_heading("Python Example", level=2)
    code_py = (
        'import fitz\n\n'
        'doc = fitz.open("input.pdf")\n'
        'for page in doc:\n'
        '    text = page.get_text()\n'
        '    print(text)'
    )
    p2 = doc.add_paragraph()
    r2 = p2.add_run(code_py)
    r2.font.name = "Courier New"
    r2.font.size = Pt(9)
    pPr2 = p2._element.get_or_add_pPr()
    shd2 = pPr2.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F5F5F5",
    })
    pPr2.append(shd2)

    doc.add_heading("JSON Example", level=2)
    code_json = (
        '{\n'
        '  "name": "MiniPdf",\n'
        '  "version": "1.0.0",\n'
        '  "features": ["xlsx", "docx", "pdf"]\n'
        '}'
    )
    p3 = doc.add_paragraph()
    r3 = p3.add_run(code_json)
    r3.font.name = "Courier New"
    r3.font.size = Pt(9)
    pPr3 = p3._element.get_or_add_pPr()
    shd3 = pPr3.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F5F5F5",
    })
    pPr3.append(shd3)
    doc.save(path)


def classic66_colored_title_page(path):
    """Title page with large colored text and decorative elements."""
    doc = Document()
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("ANNUAL REPORT 2026")
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(47, 84, 150)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Innovation Through Technology")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(89, 89, 89)
    r2.italic = True

    doc.add_paragraph()
    img_buf = _create_test_png(400, 200, (47, 84, 150))
    last_p = doc.add_paragraph()
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(img_buf, width=Inches(5))

    doc.add_paragraph()
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_org.add_run("MiniPdf Corporation")
    r3.font.size = Pt(14)
    r3.bold = True

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_date.add_run("Published: March 2026")
    r4.font.size = Pt(12)
    r4.font.color.rgb = RGBColor(128, 128, 128)
    doc.save(path)


def classic67_alternating_row_table(path):
    """Table with alternating row colors and styled headers."""
    doc = Document()
    doc.add_heading("Employee Directory", level=1)

    employees = [
        ("E001", "Alice Johnson", "Engineering", "Senior Developer", "$120,000"),
        ("E002", "Bob Williams", "Marketing", "Marketing Manager", "$95,000"),
        ("E003", "Carol Davis", "Finance", "Financial Analyst", "$88,000"),
        ("E004", "David Brown", "Engineering", "Tech Lead", "$135,000"),
        ("E005", "Emily Chen", "Design", "UX Designer", "$92,000"),
        ("E006", "Frank Miller", "Engineering", "Junior Developer", "$75,000"),
        ("E007", "Grace Lee", "HR", "HR Specialist", "$82,000"),
        ("E008", "Henry Wilson", "Engineering", "DevOps Engineer", "$110,000"),
        ("E009", "Iris Taylor", "Finance", "CFO", "$180,000"),
        ("E010", "Jack Martin", "Marketing", "Content Writer", "$68,000"),
        ("E011", "Karen White", "Engineering", "QA Engineer", "$90,000"),
        ("E012", "Leo Harris", "Design", "Graphic Designer", "$78,000"),
    ]

    table = doc.add_table(rows=len(employees) + 1, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Name", "Department", "Title", "Salary"]
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

    for ri, emp in enumerate(employees):
        row = table.rows[ri + 1]
        for ci, val in enumerate(emp):
            row.cells[ci].text = val
        if ri % 2 == 0:
            for ci in range(5):
                _set_cell_shading(row.cells[ci], "D9E2F3")
    doc.save(path)


def classic68_sidebar_layout(path):
    """Simulated sidebar layout using a 2-column table."""
    doc = Document()
    doc.add_heading("Project Documentation", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Sidebar (narrow left column)
    sidebar = table.rows[0].cells[0]
    _set_cell_shading(sidebar, "2F5496")
    sp = sidebar.paragraphs[0]
    sr = sp.add_run("Navigation")
    sr.bold = True
    sr.font.color.rgb = RGBColor(255, 255, 255)
    sr.font.size = Pt(12)
    for item in ["Overview", "Installation", "Configuration", "API Reference", "FAQ", "Changelog"]:
        p = sidebar.add_paragraph()
        r = p.add_run(item)
        r.font.color.rgb = RGBColor(200, 210, 230)
        r.font.size = Pt(10)

    # Main content (wide right column)
    main = table.rows[0].cells[1]
    mp = main.paragraphs[0]
    mr = mp.add_run("Overview")
    mr.bold = True
    mr.font.size = Pt(14)
    main.add_paragraph(
        "MiniPdf is a lightweight .NET library for converting Word and Excel "
        "documents to PDF format without requiring Microsoft Office."
    )
    mp2 = main.add_paragraph()
    mr2 = mp2.add_run("Key Features")
    mr2.bold = True
    mr2.font.size = Pt(12)
    for feat in ["DOCX to PDF conversion", "XLSX to PDF conversion",
                 "CJK font support", "Image embedding", "Table formatting"]:
        main.add_paragraph(feat, style="List Bullet")

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)
    doc.save(path)


def classic69_blockquote_styling(path):
    """Document with styled blockquotes and callout boxes."""
    doc = Document()
    doc.add_heading("Famous Quotes Collection", level=1)

    quotes = [
        ("The only way to do great work is to love what you do.", "Steve Jobs", "4472C4"),
        ("Innovation distinguishes between a leader and a follower.", "Steve Jobs", "548235"),
        ("Stay hungry, stay foolish.", "Stewart Brand", "BF8F00"),
        ("The future belongs to those who believe in the beauty of their dreams.", "Eleanor Roosevelt", "C00000"),
        ("In the middle of difficulty lies opportunity.", "Albert Einstein", "7030A0"),
    ]

    for text, author, color in quotes:
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        # Left border for blockquote effect
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        left = pBdr.makeelement(qn("w:left"), {
            qn("w:val"): "single",
            qn("w:sz"): "24",
            qn("w:space"): "4",
            qn("w:color"): color,
        })
        pBdr.append(left)
        pPr.append(pBdr)
        # Indent
        p.paragraph_format.left_indent = Inches(0.5)
        r = p.add_run(f'"{text}"')
        r.italic = True
        r.font.size = Pt(12)
        # Author
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Inches(0.5)
        ra = pa.add_run(f"  {author}")
        ra.bold = True
        ra.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()
    doc.save(path)


def classic70_academic_paper(path):
    """Academic paper with abstract, sections, and references."""
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("A Survey of Modern PDF Generation Techniques")
    r.bold = True
    r.font.size = Pt(16)

    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_authors.add_run("John Smith, Jane Doe, Robert Johnson")
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = p_inst.add_run("Department of Computer Science, University of Technology")
    ri.italic = True
    ri.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("Abstract", level=2)
    p_abs = doc.add_paragraph(
        "This paper surveys modern techniques for generating PDF documents from structured "
        "office formats. We evaluate the quality and performance of conversion tools across "
        "a benchmark suite of 90 test documents covering diverse formatting features."
    )
    p_abs.paragraph_format.left_indent = Inches(0.5)
    p_abs.paragraph_format.right_indent = Inches(0.5)

    doc.add_heading("1. Introduction", level=2)
    doc.add_paragraph(
        "PDF (Portable Document Format) remains the standard for sharing documents "
        "with consistent visual fidelity. Converting from editable office formats such "
        "as DOCX and XLSX to PDF requires careful handling of fonts, layouts, images, "
        "and styling attributes."
    )

    doc.add_heading("2. Methodology", level=2)
    doc.add_paragraph(
        "Our benchmark suite consists of 90 DOCX test files and 90 XLSX test files, "
        "each targeting specific formatting features. We compare output from MiniPdf "
        "against LibreOffice-generated reference PDFs using pixel-level similarity scoring."
    )

    doc.add_heading("3. Results", level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    for ci, h in enumerate(["Feature Category", "Avg Score", "Sample Size"]):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    data = [
        ("Basic Text", "98.5%", "20"),
        ("Tables", "96.2%", "25"),
        ("Images", "94.8%", "15"),
        ("Mixed Content", "95.1%", "30"),
    ]
    for ri, (cat, score, size) in enumerate(data):
        table.rows[ri + 1].cells[0].text = cat
        table.rows[ri + 1].cells[1].text = score
        table.rows[ri + 1].cells[2].text = size

    doc.add_heading("4. Conclusion", level=2)
    doc.add_paragraph(
        "Modern lightweight PDF generation libraries can achieve high fidelity output "
        "for the majority of common document formatting features."
    )

    doc.add_heading("References", level=2)
    refs = [
        "[1] ISO 32000-2:2020. Document management - Portable document format.",
        "[2] ECMA-376. Office Open XML File Formats.",
        "[3] Smith et al. (2025). Benchmark-driven development for document conversion.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.size = Pt(10)
    doc.save(path)


def classic71_legal_document(path):
    """Legal document with numbered clauses and sub-clauses."""
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("SOFTWARE LICENSE AGREEMENT")
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph(
        "This Software License Agreement (the \"Agreement\") is entered into as of "
        "March 1, 2026, by and between MiniPdf Corporation (\"Licensor\") and the "
        "End User (\"Licensee\")."
    )
    doc.add_paragraph()

    clauses = [
        ("1. GRANT OF LICENSE", [
            "1.1 Subject to the terms of this Agreement, Licensor grants Licensee a non-exclusive, "
            "non-transferable license to use the Software.",
            "1.2 The license granted herein is limited to use on a single computer system.",
            "1.3 Licensee may make one backup copy of the Software for archival purposes.",
        ]),
        ("2. RESTRICTIONS", [
            "2.1 Licensee shall not reverse engineer, decompile, or disassemble the Software.",
            "2.2 Licensee shall not sublicense, rent, or lease the Software to third parties.",
            "2.3 Licensee shall not remove any proprietary notices from the Software.",
        ]),
        ("3. INTELLECTUAL PROPERTY", [
            "3.1 The Software is protected by copyright and other intellectual property laws.",
            "3.2 Licensor retains all right, title, and interest in and to the Software.",
        ]),
        ("4. WARRANTY DISCLAIMER", [
            "4.1 THE SOFTWARE IS PROVIDED \"AS IS\" WITHOUT WARRANTY OF ANY KIND.",
            "4.2 LICENSOR DISCLAIMS ALL WARRANTIES, EXPRESS OR IMPLIED.",
        ]),
        ("5. LIMITATION OF LIABILITY", [
            "5.1 IN NO EVENT SHALL LICENSOR BE LIABLE FOR ANY INDIRECT, INCIDENTAL, "
            "OR CONSEQUENTIAL DAMAGES.",
            "5.2 LICENSOR'S TOTAL LIABILITY SHALL NOT EXCEED THE AMOUNT PAID BY LICENSEE.",
        ]),
    ]

    for title, subclauses in clauses:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        for sub in subclauses:
            sp = doc.add_paragraph(sub)
            sp.paragraph_format.left_indent = Inches(0.5)
            sp.paragraph_format.space_before = Pt(4)
    doc.save(path)


def classic72_technical_specification(path):
    """Technical specification document with tables and requirements."""
    doc = Document()
    doc.add_heading("Technical Specification: PDF Converter v2.0", level=1)
    doc.add_paragraph("Document Version: 2.0 | Last Updated: March 2026")
    doc.add_paragraph()

    doc.add_heading("1. System Requirements", level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    reqs = [
        ("Operating System", "Windows 10+, Linux, macOS 12+"),
        ("Runtime", ".NET 8.0 or later"),
        ("Memory", "Minimum 512 MB RAM"),
        ("Disk Space", "50 MB for installation"),
        ("Dependencies", "No external dependencies"),
    ]
    table.rows[0].cells[0].text = "Requirement"
    table.rows[0].cells[1].text = "Specification"
    _set_cell_shading(table.rows[0].cells[0], "2F5496")
    _set_cell_shading(table.rows[0].cells[1], "2F5496")
    for r in table.rows[0].cells[0].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
    for r in table.rows[0].cells[1].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
    for ri, (req, spec) in enumerate(reqs):
        table.rows[ri + 1].cells[0].text = req
        table.rows[ri + 1].cells[1].text = spec

    doc.add_heading("2. Feature Matrix", level=2)
    features = doc.add_table(rows=8, cols=3)
    features.style = "Table Grid"
    fheaders = ["Feature", "Status", "Priority"]
    for ci, h in enumerate(fheaders):
        cell = features.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True
    fdata = [
        ("DOCX to PDF", "Implemented", "High"),
        ("XLSX to PDF", "Implemented", "High"),
        ("Chart rendering", "Implemented", "Medium"),
        ("CJK font support", "Implemented", "High"),
        ("Image embedding", "Implemented", "Medium"),
        ("Hyperlink support", "Planned", "Low"),
        ("SVG support", "Planned", "Low"),
    ]
    for ri, (feat, status, pri) in enumerate(fdata):
        row = features.rows[ri + 1]
        row.cells[0].text = feat
        p = row.cells[1].paragraphs[0]
        r = p.add_run(status)
        r.font.color.rgb = RGBColor(0, 128, 0) if status == "Implemented" else RGBColor(255, 165, 0)
        row.cells[2].text = pri

    doc.add_heading("3. Performance Targets", level=2)
    doc.add_paragraph("The converter shall meet the following performance criteria:")
    for item in [
        "Convert a 10-page DOCX in under 2 seconds",
        "Convert a 100-row XLSX in under 3 seconds",
        "Memory usage shall not exceed 200 MB for standard documents",
        "Output PDF size shall be within 2x of reference PDF size",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(path)


def classic73_calendar_layout(path):
    """Monthly calendar layout using a table."""
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("March 2026")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(47, 84, 150)

    table = doc.add_table(rows=6, cols=7)
    table.style = "Table Grid"
    days = ["Sun", "Mon", "Tue", "Wed", "Thu", "Fri", "Sat"]
    for ci, d in enumerate(days):
        cell = table.rows[0].cells[ci]
        cell.text = d
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

    # March 2026 starts on Sunday
    cal = [
        [1, 2, 3, 4, 5, 6, 7],
        [8, 9, 10, 11, 12, 13, 14],
        [15, 16, 17, 18, 19, 20, 21],
        [22, 23, 24, 25, 26, 27, 28],
        [29, 30, 31, 0, 0, 0, 0],
    ]
    for ri, week in enumerate(cal):
        for ci, day in enumerate(week):
            cell = table.rows[ri + 1].cells[ci]
            if day > 0:
                cell.text = str(day)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ci == 0 or ci == 6:
                _set_cell_shading(cell, "F2F2F2")
    doc.save(path)


def classic74_org_chart(path):
    """Simple organization chart using nested tables."""
    doc = Document()
    doc.add_heading("Organization Chart", level=1)

    # CEO level
    t_top = doc.add_table(rows=1, cols=3)
    t_top.rows[0].cells[0].text = ""
    ceo_cell = t_top.rows[0].cells[1]
    ceo_p = ceo_cell.paragraphs[0]
    ceo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ceo_r = ceo_p.add_run("CEO\nJane Smith")
    ceo_r.bold = True
    _set_cell_shading(ceo_cell, "2F5496")
    ceo_r.font.color.rgb = RGBColor(255, 255, 255)
    t_top.rows[0].cells[2].text = ""

    doc.add_paragraph()

    # VP level
    t_vp = doc.add_table(rows=1, cols=3)
    t_vp.style = "Table Grid"
    vps = [
        ("VP Engineering\nBob Williams", "4472C4"),
        ("VP Marketing\nAlice Johnson", "4472C4"),
        ("VP Finance\nCarol Davis", "4472C4"),
    ]
    for ci, (name, color) in enumerate(vps):
        cell = t_vp.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
        _set_cell_shading(cell, color)

    doc.add_paragraph()

    # Team level
    t_team = doc.add_table(rows=2, cols=3)
    t_team.style = "Table Grid"
    teams = [
        [("Frontend\n3 members", "D9E2F3"), ("Brand\n2 members", "D9E2F3"), ("Accounting\n4 members", "D9E2F3")],
        [("Backend\n5 members", "D9E2F3"), ("Content\n3 members", "D9E2F3"), ("Planning\n2 members", "D9E2F3")],
    ]
    for ri, row in enumerate(teams):
        for ci, (name, color) in enumerate(row):
            cell = t_team.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(name)
            _set_cell_shading(cell, color)
    doc.save(path)


def classic75_newsletter_layout(path):
    """Newsletter with sections, images, and columns."""
    doc = Document()
    # Masthead
    p_mast = doc.add_paragraph()
    p_mast.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_mast.add_run("THE TECH WEEKLY")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(47, 84, 150)
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = p_date.add_run("Issue #42 | March 6, 2026")
    rd.font.size = Pt(10)
    rd.font.color.rgb = RGBColor(128, 128, 128)
    _add_horizontal_rule(doc)

    # Lead story
    doc.add_heading("AI Revolution in Software Development", level=2)
    img_buf = _create_test_png(300, 150, (30, 90, 180))
    doc.add_picture(img_buf, width=Inches(4))
    doc.add_paragraph(
        "The integration of large language models into development workflows is "
        "transforming how teams write, review, and deploy code. Industry experts "
        "predict that AI-assisted development will become standard practice by 2027."
    )

    _add_horizontal_rule(doc)

    # Two-column news
    doc.add_heading("Quick Updates", level=2)
    news_table = doc.add_table(rows=1, cols=2)
    left = news_table.rows[0].cells[0]
    lp = left.paragraphs[0]
    lr = lp.add_run("Cloud Infrastructure")
    lr.bold = True
    left.add_paragraph(
        "Major cloud providers announce new edge computing regions in Asia-Pacific."
    )
    right = news_table.rows[0].cells[1]
    rp = right.paragraphs[0]
    rr = rp.add_run("Open Source")
    rr.bold = True
    right.add_paragraph(
        "The .NET Foundation releases new guidelines for community project governance."
    )
    doc.save(path)


def classic76_recipe_card(path):
    """Recipe card document with ingredients and instructions."""
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Classic Chocolate Chip Cookies")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(139, 69, 19)

    # Info table
    info = doc.add_table(rows=1, cols=3)
    info.style = "Table Grid"
    for ci, (label, val) in enumerate([
        ("Prep Time", "15 min"),
        ("Cook Time", "12 min"),
        ("Servings", "36 cookies"),
    ]):
        cell = info.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}\n")
        rl.bold = True
        rl.font.size = Pt(9)
        rv = p.add_run(val)
        rv.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_heading("Ingredients", level=2)
    ingredients = [
        "2 1/4 cups all-purpose flour",
        "1 tsp baking soda",
        "1 tsp salt",
        "1 cup (2 sticks) butter, softened",
        "3/4 cup granulated sugar",
        "3/4 cup packed brown sugar",
        "2 large eggs",
        "2 tsp vanilla extract",
        "2 cups chocolate chips",
    ]
    for ing in ingredients:
        doc.add_paragraph(ing, style="List Bullet")

    doc.add_heading("Instructions", level=2)
    steps = [
        "Preheat oven to 375 degrees F.",
        "Combine flour, baking soda, and salt in a small bowl.",
        "Beat butter, granulated sugar, and brown sugar in a large mixer bowl until creamy.",
        "Add eggs and vanilla extract; beat well.",
        "Gradually blend in flour mixture.",
        "Stir in chocolate chips.",
        "Drop rounded tablespoon of dough onto ungreased baking sheets.",
        "Bake for 9 to 11 minutes or until golden brown.",
        "Cool on baking sheets for 2 minutes; remove to wire racks to cool completely.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.save(path)


def classic77_timeline_layout(path):
    """Project timeline using a table layout."""
    doc = Document()
    doc.add_heading("Project Timeline: MiniPdf v2.0", level=1)

    milestones = [
        ("Q1 2025", "Project Inception", "Initial requirements gathering and architecture design.", "4472C4"),
        ("Q2 2025", "XLSX Support", "Implemented Excel-to-PDF conversion with chart support.", "548235"),
        ("Q3 2025", "Benchmark Suite", "Created automated benchmark pipeline with 60 test cases.", "BF8F00"),
        ("Q4 2025", "CJK Support", "Added Chinese, Japanese, and Korean font embedding.", "C00000"),
        ("Q1 2026", "DOCX Support", "Implemented Word-to-PDF conversion achieving 97% quality.", "7030A0"),
        ("Q2 2026", "v2.0 Release", "Public release with full documentation and NuGet package.", "2F5496"),
    ]

    table = doc.add_table(rows=len(milestones), cols=3)
    table.style = "Table Grid"
    for ri, (date, title, desc, color) in enumerate(milestones):
        # Date column
        dc = table.rows[ri].cells[0]
        dc.width = Inches(1)
        dp = dc.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = dp.add_run(date)
        dr.bold = True
        dr.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(dc, color)

        # Title column
        tc = table.rows[ri].cells[1]
        tc.width = Inches(1.5)
        tp = tc.paragraphs[0]
        tr = tp.add_run(title)
        tr.bold = True

        # Description column
        table.rows[ri].cells[2].text = desc
    doc.save(path)


def classic78_faq_document(path):
    """FAQ document with questions and answers."""
    doc = Document()
    doc.add_heading("Frequently Asked Questions", level=1)
    doc.add_paragraph("Find answers to the most common questions about MiniPdf.")
    doc.add_paragraph()

    faqs = [
        ("What is MiniPdf?",
         "MiniPdf is a lightweight .NET library for converting DOCX and XLSX files to PDF "
         "without requiring Microsoft Office installation. It supports both DOCX and XLSX formats."),
        ("Which .NET versions are supported?",
         ".NET 8.0 and later versions are supported. The library targets .NET Standard 2.0 "
         "for maximum compatibility across different .NET implementations."),
        ("Does MiniPdf support images in documents?",
         "Yes, MiniPdf supports embedded images in both DOCX and XLSX formats. Images are "
         "converted and embedded in the output PDF with proper scaling."),
        ("How is the conversion quality measured?",
         "Quality is measured by pixel-level comparison against LibreOffice-generated reference "
         "PDFs. Each test case receives a similarity score from 0% to 100%."),
        ("Can MiniPdf handle CJK characters?",
         "Yes, MiniPdf includes CJK font embedding support for Chinese, Japanese, and Korean "
         "text in both DOCX and XLSX documents."),
        ("Is MiniPdf available on NuGet?",
         "Yes, MiniPdf is published as a NuGet package and can be installed via "
         "dotnet add package MiniPdf."),
        ("What table features are supported?",
         "MiniPdf supports table borders, cell shading, merged cells, column widths, and "
         "alternating row colors in both DOCX and XLSX formats."),
        ("How do I report a bug?",
         "Please open an issue on the GitHub repository with a minimal reproduction case "
         "and the expected vs actual output."),
        ("Does MiniPdf support headers and footers?",
         "MiniPdf does not currently render headers and footers from DOCX files. The content "
         "area of each page is fully supported including text, tables, and images."),
        ("What is the maximum file size supported?",
         "There is no hard limit on file size. MiniPdf processes files in a streaming fashion "
         "and memory usage scales with document complexity rather than file size."),
        ("Can I use MiniPdf in a web application?",
         "Yes, MiniPdf works in any .NET environment including ASP.NET Core web applications, "
         "Azure Functions, and containerized services."),
        ("Does MiniPdf preserve hyperlinks?",
         "Hyperlink text is preserved in the output PDF, but clickable URLs are not currently "
         "supported. The link text appears as regular styled text."),
    ]

    for i, (q, a) in enumerate(faqs, 1):
        pq = doc.add_paragraph()
        rq = pq.add_run(f"Q{i}: {q}")
        rq.bold = True
        rq.font.color.rgb = RGBColor(47, 84, 150)

        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Inches(0.3)
        ra = pa.add_run(f"A: {a}")
        doc.add_paragraph()
    doc.save(path)


def classic79_glossary(path):
    """Glossary / definition list document."""
    doc = Document()
    doc.add_heading("Glossary of Terms", level=1)

    terms = [
        ("API", "Application Programming Interface. A set of protocols and tools for building software applications."),
        ("CI/CD", "Continuous Integration / Continuous Deployment. Practices for automating software delivery."),
        ("CJK", "Chinese, Japanese, Korean. Refers to the character sets used in these languages."),
        ("DOCX", "The XML-based file format for Microsoft Word documents, defined by ECMA-376."),
        ("EMU", "English Metric Unit. The base unit of measurement in OOXML documents (1 inch = 914400 EMU)."),
        ("NuGet", "The package manager for .NET, used to distribute and consume .NET libraries."),
        ("OOXML", "Office Open XML. The ISO-standardized format used by Microsoft Office."),
        ("PDF", "Portable Document Format. An ISO standard for document exchange (ISO 32000)."),
        ("SSIM", "Structural Similarity Index Measure. A metric for predicting image quality."),
        ("XLSX", "The XML-based file format for Microsoft Excel workbooks, defined by ECMA-376."),
    ]

    for term, definition in terms:
        pt = doc.add_paragraph()
        rt = pt.add_run(term)
        rt.bold = True
        rt.font.size = Pt(12)
        rt.font.color.rgb = RGBColor(47, 84, 150)

        pd = doc.add_paragraph(definition)
        pd.paragraph_format.left_indent = Inches(0.5)
        pd.paragraph_format.space_after = Pt(12)
    doc.save(path)


def classic80_matrix_grid(path):
    """Responsibility assignment matrix (RACI chart)."""
    doc = Document()
    doc.add_heading("RACI Matrix - MiniPdf Project", level=1)

    tasks = [
        "Requirements gathering",
        "Architecture design",
        "DOCX parser implementation",
        "XLSX parser implementation",
        "PDF writer implementation",
        "Benchmark suite creation",
        "CJK font support",
        "Code review",
        "Documentation",
        "Release management",
    ]
    roles = ["Project Lead", "Dev Team", "QA", "DevOps"]
    raci_data = [
        ["A", "R", "C", "I"],
        ["R", "C", "I", "C"],
        ["A", "R", "C", "I"],
        ["A", "R", "C", "I"],
        ["A", "R", "I", "I"],
        ["C", "R", "A", "I"],
        ["A", "R", "C", "I"],
        ["C", "R", "A", "I"],
        ["A", "R", "C", "C"],
        ["R", "I", "C", "A"],
    ]
    colors_map = {
        "R": ("0070C0", "DAEEF3"),
        "A": ("C00000", "FCE4EC"),
        "C": ("548235", "E2EFDA"),
        "I": ("7F7F7F", "F2F2F2"),
    }

    table = doc.add_table(rows=len(tasks) + 1, cols=len(roles) + 1)
    table.style = "Table Grid"
    # Header row
    table.rows[0].cells[0].text = "Task"
    _set_cell_shading(table.rows[0].cells[0], "2F5496")
    for r in table.rows[0].cells[0].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.bold = True
    for ci, role in enumerate(roles):
        cell = table.rows[0].cells[ci + 1]
        cell.text = role
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

    for ri, task in enumerate(tasks):
        row = table.rows[ri + 1]
        row.cells[0].text = task
        for ci, val in enumerate(raci_data[ri]):
            cell = row.cells[ci + 1]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fg, bg = colors_map[val]
            rr = p.add_run(val)
            rr.bold = True
            rr.font.color.rgb = RGBColor(
                int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16)
            )
            _set_cell_shading(cell, bg)
    doc.save(path)


def classic81_budget_table(path):
    """Budget document with financial table and totals."""
    doc = Document()
    doc.add_heading("Annual Budget Report - FY2026", level=1)

    categories = [
        ("Personnel", [("Salaries", 450000), ("Benefits", 135000), ("Training", 25000)]),
        ("Infrastructure", [("Cloud Services", 120000), ("Hardware", 45000), ("Licenses", 32000)]),
        ("Operations", [("Office Rent", 96000), ("Utilities", 18000), ("Supplies", 8000)]),
        ("Marketing", [("Digital Ads", 60000), ("Events", 35000), ("Content", 20000)]),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for ci, h in enumerate(["Category / Item", "Budget ($)", "% of Total"]):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

    grand_total = sum(amt for _, items in categories for _, amt in items)

    for cat_name, items in categories:
        cat_total = sum(amt for _, amt in items)
        # Category header row
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        rr = p.add_run(cat_name)
        rr.bold = True
        row.cells[1].text = f"{cat_total:,.0f}"
        for r in row.cells[1].paragraphs[0].runs:
            r.bold = True
        row.cells[2].text = f"{cat_total / grand_total * 100:.1f}%"
        for r in row.cells[2].paragraphs[0].runs:
            r.bold = True
        _set_cell_shading(row.cells[0], "D9E2F3")
        _set_cell_shading(row.cells[1], "D9E2F3")
        _set_cell_shading(row.cells[2], "D9E2F3")

        for item_name, amt in items:
            irow = table.add_row()
            irow.cells[0].text = f"    {item_name}"
            irow.cells[1].text = f"{amt:,.0f}"
            irow.cells[2].text = f"{amt / grand_total * 100:.1f}%"

    # Grand total row
    trow = table.add_row()
    p = trow.cells[0].paragraphs[0]
    rr = p.add_run("GRAND TOTAL")
    rr.bold = True
    trow.cells[1].text = f"{grand_total:,.0f}"
    for r in trow.cells[1].paragraphs[0].runs:
        r.bold = True
    trow.cells[2].text = "100.0%"
    for r in trow.cells[2].paragraphs[0].runs:
        r.bold = True
    _set_cell_shading(trow.cells[0], "2F5496")
    _set_cell_shading(trow.cells[1], "2F5496")
    _set_cell_shading(trow.cells[2], "2F5496")
    for ci in range(3):
        for r in trow.cells[ci].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
    doc.save(path)


def classic82_survey_questionnaire(path):
    """Survey/questionnaire form document."""
    doc = Document()
    doc.add_heading("Employee Satisfaction Survey", level=1)
    doc.add_paragraph("Please rate each item on a scale of 1 (Strongly Disagree) to 5 (Strongly Agree).")
    doc.add_paragraph()

    sections = [
        ("Work Environment", [
            "My workspace is comfortable and well-equipped.",
            "The office environment supports productivity.",
            "I have the tools I need to do my job effectively.",
        ]),
        ("Management", [
            "My manager provides clear direction and expectations.",
            "I receive regular and constructive feedback.",
            "Management is transparent about company goals.",
        ]),
        ("Career Development", [
            "I have opportunities for professional growth.",
            "Training programs are relevant and accessible.",
            "There is a clear path for career advancement.",
        ]),
        ("Work-Life Balance", [
            "I can maintain a healthy work-life balance.",
            "Flexible work arrangements are available.",
            "Workload is reasonable and manageable.",
        ]),
    ]

    for section_title, questions in sections:
        doc.add_heading(section_title, level=2)
        table = doc.add_table(rows=len(questions) + 1, cols=6)
        table.style = "Table Grid"
        headers = ["Statement", "1", "2", "3", "4", "5"]
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = h
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "4472C4")
            for r in cell.paragraphs[0].runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
                r.font.size = Pt(9)
        for qi, q in enumerate(questions):
            row = table.rows[qi + 1]
            row.cells[0].text = q
            for ci in range(1, 6):
                row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    doc.save(path)


def classic83_medical_form(path):
    """Medical intake form with patient information fields."""
    doc = Document()
    doc.add_heading("Patient Intake Form", level=1)
    doc.add_paragraph("Please complete all sections. All information is kept confidential.")
    doc.add_paragraph()

    doc.add_heading("Personal Information", level=2)
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = "Table Grid"
    fields = [
        [("First Name:", ""), ("Last Name:", "")],
        [("Date of Birth:", ""), ("Gender:", "")],
        [("Phone:", ""), ("Email:", "")],
        [("Address:", ""), ("City/State/ZIP:", "")],
    ]
    for ri, row_fields in enumerate(fields):
        for fi, (label, _) in enumerate(row_fields):
            label_ci = fi * 2
            cell = info_table.rows[ri].cells[label_ci]
            p = cell.paragraphs[0]
            rr = p.add_run(label)
            rr.bold = True
            rr.font.size = Pt(9)
            _set_cell_shading(cell, "F2F2F2")

    doc.add_paragraph()
    doc.add_heading("Medical History", level=2)
    conditions = [
        "Heart Disease", "Diabetes", "High Blood Pressure", "Asthma",
        "Allergies", "Cancer", "Thyroid Disorder", "Arthritis",
    ]
    med_table = doc.add_table(rows=len(conditions), cols=3)
    med_table.style = "Table Grid"
    for ri, cond in enumerate(conditions):
        med_table.rows[ri].cells[0].text = cond
        med_table.rows[ri].cells[1].text = "Yes / No"
        med_table.rows[ri].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        med_table.rows[ri].cells[2].text = "Notes:"

    doc.add_paragraph()
    doc.add_heading("Current Medications", level=2)
    med_list_table = doc.add_table(rows=5, cols=4)
    med_list_table.style = "Table Grid"
    for ci, h in enumerate(["Medication Name", "Dosage", "Frequency", "Purpose"]):
        cell = med_list_table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "4472C4")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True
            r.font.size = Pt(9)
    doc.save(path)


def classic84_shipping_label(path):
    """Shipping label document with sender/receiver info."""
    doc = Document()

    for label_num in range(1, 4):
        if label_num > 1:
            doc.add_paragraph()
            _add_horizontal_rule(doc)
            doc.add_paragraph()

        # Outer table for label
        outer = doc.add_table(rows=3, cols=2)
        outer.style = "Table Grid"

        # From
        from_cell = outer.rows[0].cells[0]
        fp = from_cell.paragraphs[0]
        fr = fp.add_run("FROM:")
        fr.bold = True
        fr.font.size = Pt(8)
        from_cell.add_paragraph(f"MiniPdf Corp\n123 Tech Ave\nSuite {100 + label_num}\nSan Francisco, CA 94105")

        # Tracking
        track_cell = outer.rows[0].cells[1]
        tp = track_cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tr = tp.add_run(f"TRACKING #:\n1Z999AA{label_num}0123456784")
        tr.font.name = "Courier New"
        tr.font.size = Pt(9)

        # To (merged row)
        to_cell = outer.rows[1].cells[0].merge(outer.rows[1].cells[1])
        top = to_cell.paragraphs[0]
        tor = top.add_run("TO:")
        tor.bold = True
        tor.font.size = Pt(8)
        addr = to_cell.add_paragraph()
        ar = addr.add_run(f"Customer {label_num}\n{456 + label_num * 100} Main Street\nNew York, NY 1000{label_num}")
        ar.font.size = Pt(14)
        ar.bold = True

        # Shipping info
        info_cell = outer.rows[2].cells[0].merge(outer.rows[2].cells[1])
        ip = info_cell.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ir = ip.add_run(f"PRIORITY MAIL | Weight: {label_num * 2.5:.1f} lbs | Ship Date: 03/0{label_num}/2026")
        ir.font.size = Pt(9)
        _set_cell_shading(info_cell, "F2F2F2")
    doc.save(path)


def classic85_report_card(path):
    """Student report card / transcript document."""
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("STUDENT REPORT CARD")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(47, 84, 150)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.add_run("Academic Year 2025-2026 | Spring Semester")

    doc.add_paragraph()

    # Student info
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    info_data = [
        [("Student:", "John Smith"), ("Grade:", "10th")],
        [("Student ID:", "STU-2026-001"), ("Advisor:", "Ms. Johnson")],
    ]
    for ri, row in enumerate(info_data):
        for fi, (label, val) in enumerate(row):
            lc = info.rows[ri].cells[fi * 2]
            lp = lc.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(9)
            _set_cell_shading(lc, "F2F2F2")
            vc = info.rows[ri].cells[fi * 2 + 1]
            vc.text = val

    doc.add_paragraph()

    # Grades table
    grades_table = doc.add_table(rows=9, cols=5)
    grades_table.style = "Table Grid"
    g_headers = ["Subject", "Teacher", "Grade", "Score", "Credits"]
    for ci, h in enumerate(g_headers):
        cell = grades_table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

    subjects = [
        ("Mathematics", "Mr. Thompson", "A", "95", "4"),
        ("English", "Ms. Williams", "A-", "91", "4"),
        ("Physics", "Dr. Brown", "B+", "88", "4"),
        ("History", "Ms. Davis", "A", "94", "3"),
        ("Computer Science", "Mr. Lee", "A+", "98", "3"),
        ("Art", "Ms. Garcia", "B+", "87", "2"),
        ("Physical Education", "Coach Miller", "A", "96", "1"),
        ("GPA", "", "3.78", "", "21"),
    ]
    for ri, vals in enumerate(subjects):
        row = grades_table.rows[ri + 1]
        for ci, val in enumerate(vals):
            row.cells[ci].text = val
        if ri == len(subjects) - 1:  # GPA row
            for ci in range(5):
                _set_cell_shading(row.cells[ci], "D9E2F3")
                for r in row.cells[ci].paragraphs[0].runs:
                    r.bold = True
    doc.save(path)


def classic86_checklist_document(path):
    """Checklist document with checkboxes (using symbols)."""
    doc = Document()
    doc.add_heading("Project Launch Checklist", level=1)
    doc.add_paragraph("Complete all items before the scheduled launch date.")
    doc.add_paragraph()

    checklists = [
        ("Pre-Launch", [
            (True, "Code review completed"),
            (True, "Unit tests passing"),
            (True, "Integration tests passing"),
            (False, "Performance benchmarks met"),
            (True, "Security audit completed"),
        ]),
        ("Documentation", [
            (True, "API documentation updated"),
            (True, "README file updated"),
            (False, "Changelog written"),
            (False, "Migration guide prepared"),
        ]),
        ("Deployment", [
            (False, "Staging environment tested"),
            (False, "Production config verified"),
            (False, "Rollback plan documented"),
            (False, "Monitoring alerts configured"),
            (False, "DNS records updated"),
        ]),
        ("Post-Launch", [
            (False, "Announce on social media"),
            (False, "Send newsletter"),
            (False, "Update project website"),
            (False, "Close related GitHub issues"),
        ]),
    ]

    for section, items in checklists:
        doc.add_heading(section, level=2)
        for checked, text in items:
            p = doc.add_paragraph()
            symbol = "\u2611" if checked else "\u2610"  # checked/unchecked box
            r = p.add_run(f"  {symbol}  {text}")
            if checked:
                r.font.color.rgb = RGBColor(0, 128, 0)
    doc.save(path)


def classic87_bibliography(path):
    """Bibliography / references page with formatted citations."""
    doc = Document()
    doc.add_heading("Bibliography", level=1)
    doc.add_paragraph()

    entries = [
        {
            "authors": "Smith, J., & Johnson, R.",
            "year": "2025",
            "title": "Modern Document Processing: A Comprehensive Survey",
            "journal": "Journal of Software Engineering",
            "volume": "42(3), 125-148",
        },
        {
            "authors": "Chen, L., Wang, M., & Liu, X.",
            "year": "2024",
            "title": "Benchmarking PDF Generation Libraries: Quality and Performance",
            "journal": "ACM Computing Surveys",
            "volume": "57(1), 1-35",
        },
        {
            "authors": "Brown, A.",
            "year": "2025",
            "title": "Office Open XML: Architecture and Implementation",
            "journal": "IEEE Transactions on Document Analysis",
            "volume": "28(4), 890-905",
        },
        {
            "authors": "ISO",
            "year": "2020",
            "title": "ISO 32000-2:2020 Document Management - Portable Document Format",
            "journal": "International Organization for Standardization",
            "volume": "",
        },
        {
            "authors": "ECMA International",
            "year": "2016",
            "title": "ECMA-376: Office Open XML File Formats",
            "journal": "ECMA International",
            "volume": "5th Edition",
        },
        {
            "authors": "Davis, K., & Martinez, S.",
            "year": "2024",
            "title": "AI-Driven Code Review: Patterns and Anti-Patterns",
            "journal": "Proceedings of ICSE 2024",
            "volume": "pp. 445-460",
        },
        {
            "authors": "Taylor, P.",
            "year": "2025",
            "title": "CJK Font Embedding in Portable Documents",
            "journal": "International Journal of Digital Typography",
            "volume": "15(2), 78-92",
        },
    ]

    for i, entry in enumerate(entries, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

        r_num = p.add_run(f"[{i}] ")
        r_num.bold = True
        p.add_run(f"{entry['authors']} ({entry['year']}). ")
        r_title = p.add_run(f"{entry['title']}. ")
        r_title.italic = True
        r_journal = p.add_run(entry['journal'])
        if entry['volume']:
            p.add_run(f", {entry['volume']}.")
        else:
            p.add_run(".")
    doc.save(path)


def classic88_presentation_handout(path):
    """Presentation handout with slides and notes areas."""
    doc = Document()
    doc.add_heading("Presentation Handout", level=1)
    doc.add_paragraph("MiniPdf: Lightweight Document Conversion for .NET")
    doc.add_paragraph()

    slides = [
        ("Slide 1: Introduction", "What is MiniPdf and why does it matter?",
         (47, 84, 150)),
        ("Slide 2: Architecture", "Core components and design decisions.",
         (84, 130, 53)),
        ("Slide 3: DOCX Support", "How Word documents are parsed and converted.",
         (191, 144, 0)),
        ("Slide 4: XLSX Support", "Excel workbook processing and chart rendering.",
         (192, 0, 0)),
        ("Slide 5: Quality Assurance", "Benchmark pipeline and self-evolution cycle.",
         (112, 48, 160)),
        ("Slide 6: Future Plans", "Roadmap for v2.0 and beyond.",
         (0, 112, 192)),
    ]

    for title, description, color in slides:
        # Slide placeholder
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"

        # Slide image area
        slide_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
        _set_cell_shading(slide_cell, f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
        sp = slide_cell.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(f"\n{title}\n")
        sr.bold = True
        sr.font.size = Pt(16)
        sr.font.color.rgb = RGBColor(255, 255, 255)
        sd = slide_cell.add_paragraph()
        sd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sdr = sd.add_run(description)
        sdr.font.color.rgb = RGBColor(220, 220, 220)

        # Notes area
        notes_cell = table.rows[1].cells[0].merge(table.rows[1].cells[1])
        np = notes_cell.paragraphs[0]
        nr = np.add_run("Notes:")
        nr.bold = True
        nr.font.size = Pt(9)
        nr.font.color.rgb = RGBColor(128, 128, 128)
        for _ in range(3):
            notes_cell.add_paragraph("_" * 70)

        doc.add_paragraph()
    doc.save(path)


def classic89_multi_image_gallery(path):
    """Gallery layout with multiple images in a grid."""
    doc = Document()
    doc.add_heading("Image Gallery", level=1)
    doc.add_paragraph("Collection of test images for benchmark validation.")
    doc.add_paragraph()

    colors = [
        ((70, 130, 180), "Steel Blue"),
        ((220, 20, 60), "Crimson"),
        ((50, 205, 50), "Lime Green"),
        ((255, 165, 0), "Orange"),
        ((147, 112, 219), "Purple"),
        ((0, 206, 209), "Turquoise"),
        ((255, 215, 0), "Gold"),
        ((188, 143, 143), "Rosy Brown"),
        ((100, 149, 237), "Cornflower"),
    ]

    # 3x3 grid
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    idx = 0
    for ri in range(3):
        for ci in range(3):
            color, name = colors[idx]
            cell = table.rows[ri].cells[ci]
            img_buf = _create_test_png(120, 80, color)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_buf, width=Inches(1.5))
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(name)
            cr.font.size = Pt(9)
            cr.italic = True
            idx += 1

    doc.add_paragraph()
    doc.add_paragraph("Each image demonstrates a different color channel for testing "
                      "image encoding fidelity in PDF output.")
    doc.save(path)


def classic90_comprehensive_annual_report(path):
    """Comprehensive annual report combining many features."""
    doc = Document()

    # Title page
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("MINIPDF CORPORATION")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(47, 84, 150)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Annual Report 2025-2026")
    r2.font.size = Pt(18)

    img_buf = _create_test_png(400, 150, (47, 84, 150))
    doc.add_picture(img_buf, width=Inches(5))

    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    r3 = p_org.add_run("Published: March 2026")
    r3.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Table of contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("Executive Summary", "2"),
        ("Financial Highlights", "3"),
        ("Product Development", "5"),
        ("Market Analysis", "7"),
        ("Team & Organization", "9"),
        ("Outlook & Strategy", "10"),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(title)
        p.add_run(f"{'.' * (50 - len(title))}{page}")

    doc.add_page_break()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    p_exec = doc.add_paragraph()
    p_exec.add_run("Dear Stakeholders, ").bold = True
    p_exec.add_run(
        "We are pleased to present the annual report for MiniPdf Corporation. "
        "This year marked significant milestones in our product development, "
        "including the launch of DOCX-to-PDF conversion and expansion of our "
        "benchmark suite to 180 test cases."
    )

    # Key Metrics
    doc.add_heading("Key Metrics", level=2)
    metrics_table = doc.add_table(rows=2, cols=4)
    metrics_table.style = "Table Grid"
    metrics = [
        ("Revenue", "$2.4M", "+45%", "548235"),
        ("Users", "12,500", "+120%", "4472C4"),
        ("Test Cases", "180", "+200%", "BF8F00"),
        ("Quality Score", "97.2%", "+5.1%", "7030A0"),
    ]
    for ci, (label, value, change, color) in enumerate(metrics):
        top = metrics_table.rows[0].cells[ci]
        top_p = top.paragraphs[0]
        top_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(top, color)
        tr = top_p.add_run(f"{label}\n{value}")
        tr.bold = True
        tr.font.color.rgb = RGBColor(255, 255, 255)
        tr.font.size = Pt(14)

        bot = metrics_table.rows[1].cells[ci]
        bot_p = bot.paragraphs[0]
        bot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bot_p.add_run(change)
        br.font.color.rgb = RGBColor(0, 128, 0)
        br.bold = True

    doc.add_paragraph()

    # Financial Highlights
    doc.add_heading("Financial Highlights", level=1)
    fin_table = doc.add_table(rows=6, cols=4)
    fin_table.style = "Table Grid"
    fin_headers = ["", "FY2024", "FY2025", "FY2026"]
    for ci, h in enumerate(fin_headers):
        cell = fin_table.rows[0].cells[ci]
        cell.text = h
        _set_cell_shading(cell, "2F5496")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

    fin_data = [
        ("Revenue", "$800K", "$1.6M", "$2.4M"),
        ("Expenses", "$650K", "$1.1M", "$1.5M"),
        ("Net Income", "$150K", "$500K", "$900K"),
        ("R&D Investment", "$200K", "$400K", "$600K"),
        ("Headcount", "8", "15", "25"),
    ]
    for ri, row_data in enumerate(fin_data):
        for ci, val in enumerate(row_data):
            fin_table.rows[ri + 1].cells[ci].text = val
        if ri == 2:  # Net income row highlighted
            for ci in range(4):
                _set_cell_shading(fin_table.rows[ri + 1].cells[ci], "E2EFDA")
                for r in fin_table.rows[ri + 1].cells[ci].paragraphs[0].runs:
                    r.bold = True

    # Product Development
    doc.add_heading("Product Development", level=1)
    doc.add_heading("Milestones Achieved", level=2)
    for item in [
        "XLSX-to-PDF conversion with chart support",
        "DOCX-to-PDF conversion achieving 97% quality score",
        "CJK font embedding for Chinese, Japanese, and Korean",
        "Automated benchmark pipeline with 180 test cases",
        "AI-powered code review integration",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Quality Metrics", level=2)
    img_chart = _create_test_png(300, 150, (46, 84, 150))
    doc.add_picture(img_chart, width=Inches(4))
    doc.add_paragraph()

    # Recommendations
    doc.add_heading("Outlook & Strategy", level=1)
    for rec in [
        "Expand format support to include PPTX",
        "Achieve 99% average quality score",
        "Release v2.0 on NuGet with full documentation",
        "Build enterprise partnerships",
        "Establish community contributor program",
    ]:
        doc.add_paragraph(rec, style="List Number")

    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = p_footer.add_run("--- End of Annual Report ---")
    rf.italic = True
    rf.font.color.rgb = RGBColor(128, 128, 128)
    doc.save(path)


# ── Helpers ──────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color):
    """Set a table cell's background shading."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shd)


def _add_horizontal_rule(doc):
    """Add a paragraph that looks like a horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "auto",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Main ─────────────────────────────────────────────────────────────────

ALL_GENERATORS = [
    classic01_single_paragraph,
    classic02_multiple_paragraphs,
    classic03_headings,
    classic04_bold_italic,
    classic05_font_sizes,
    classic06_font_colors,
    classic07_alignment,
    classic08_bullet_list,
    classic09_numbered_list,
    classic10_simple_table,
    classic11_table_with_shading,
    classic12_merged_cells_table,
    classic13_long_document,
    classic14_mixed_content,
    classic15_indentation,
    classic16_line_spacing,
    classic17_page_break,
    classic18_embedded_image,
    classic19_multiple_images,
    classic20_table_with_many_rows,
    classic21_nested_lists,
    classic22_horizontal_rule,
    classic23_mixed_formatting_runs,
    classic24_two_column_table_layout,
    classic25_title_and_subtitle,
    classic26_table_alignment,
    classic27_long_paragraph_wrapping,
    classic28_special_characters,
    classic29_table_with_image,
    classic30_comprehensive_report,
    classic31_strikethrough_text,
    classic32_superscript_subscript,
    classic33_highlighted_text,
    classic34_paragraph_borders,
    classic35_tab_stops,
    classic36_wide_table,
    classic37_nested_table,
    classic38_table_column_widths,
    classic39_financial_report,
    classic40_resume,
    classic41_business_letter,
    classic42_meeting_minutes,
    classic43_invoice_document,
    classic44_memo,
    classic45_project_plan,
    classic46_comparison_table,
    classic47_data_dictionary,
    classic48_multi_level_headings,
    classic49_cjk_document,
    classic50_long_table_with_formatting,
    classic51_underline_styles,
    classic52_spacing_before_after,
    classic53_table_merged_complex,
    classic54_multi_font_family,
    classic55_background_shading_paragraph,
    classic56_images_and_tables_mixed,
    classic57_right_to_left_text,
    classic58_dense_paragraph_document,
    classic59_numbered_and_bullet_mixed,
    classic60_comprehensive_styled_report,
    classic61_header_and_footer,
    classic62_footnote_references,
    classic63_toc_style_headings,
    classic64_multi_column_layout,
    classic65_code_block_styling,
    classic66_colored_title_page,
    classic67_alternating_row_table,
    classic68_sidebar_layout,
    classic69_blockquote_styling,
    classic70_academic_paper,
    classic71_legal_document,
    classic72_technical_specification,
    classic73_calendar_layout,
    classic74_org_chart,
    classic75_newsletter_layout,
    classic76_recipe_card,
    classic77_timeline_layout,
    classic78_faq_document,
    classic79_glossary,
    classic80_matrix_grid,
    classic81_budget_table,
    classic82_survey_questionnaire,
    classic83_medical_form,
    classic84_shipping_label,
    classic85_report_card,
    classic86_checklist_document,
    classic87_bibliography,
    classic88_presentation_handout,
    classic89_multi_image_gallery,
    classic90_comprehensive_annual_report,
]


def main():
    parser = argparse.ArgumentParser(description="Generate classic DOCX test files")
    parser.add_argument("--outdir", default=str(OUTPUT_DIR), help="Output directory")
    args = parser.parse_args()

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    print(f"Generating {len(ALL_GENERATORS)} classic DOCX files to {outdir}/")
    print()

    passed = 0
    failed = 0
    for i, gen_func in enumerate(ALL_GENERATORS, 1):
        name = gen_func.__name__
        filename = f"docx_{name}.docx"
        filepath = outdir / filename
        try:
            gen_func(str(filepath))
            size_kb = filepath.stat().st_size / 1024
            print(f"  OK  {filename} ({size_kb:.1f} KB)")
            passed += 1
        except Exception as e:
            print(f"  ERR {filename}: {e}")
            failed += 1

    print(f"\nDone! Passed: {passed}, Failed: {failed}, Total: {len(ALL_GENERATORS)}")


if __name__ == "__main__":
    main()
